from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "guides"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "DaemonCore-FieldOps-Production-Cheat-Sheet-v6.4.0.docx"

RED = "D71920"
BLACK = "111111"
DARK = "262626"
MID = "666666"
LIGHT = "F2F2F2"
PALE = "FAF4F4"
BORDER = "D9D9D9"
WHITE = "FFFFFF"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    fill = props.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        props.append(fill)
    fill.set(qn("w:fill"), color)


def margins(cell, top=100, start=120, bottom=100, end=120):
    props = cell._tc.get_or_add_tcPr()
    tc_mar = props.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        props.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table):
    props = table._tbl.tblPr
    node = props.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        props.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:color"), BORDER)
        node.append(line)


def set_repeat_header(row):
    props = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    props.append(repeat)


def font(run, name="Aptos", size=9, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    font(run, **kwargs)
    return run


def label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, text.upper(), name="Aptos Display", size=7.5, bold=True, color=RED)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, text, name="Aptos Display", size=17 if level == 1 else 11.5, bold=True, color=BLACK)
    return p


def intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.12
    add_text(p, text, size=9.5, color=DARK)


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.2)
    p.paragraph_format.first_line_indent = Inches(-.12)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        add_text(p, bold_lead, size=8.8, bold=True)
    add_text(p, text, size=8.8, color=DARK)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    header = t.rows[0]
    set_repeat_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, DARK)
        margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, text.upper(), size=7.3, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = t.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            margins(cell, top=90, bottom=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, str(value), size=7.9, bold=idx == 0, color=BLACK if idx == 0 else DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def page_break(doc):
    doc.add_page_break()


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.top_margin = Inches(.55)
section.bottom_margin = Inches(.55)
section.left_margin = Inches(.62)
section.right_margin = Inches(.62)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.color.rgb = RGBColor.from_string(BLACK)
styles["Title"].font.bold = True
title_style_border = styles["Title"].element.get_or_add_pPr().find(qn("w:pBdr"))
if title_style_border is not None:
    styles["Title"].element.get_or_add_pPr().remove(title_style_border)
for style_name in ("Heading 1", "Heading 2"):
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor.from_string(BLACK)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_text(hp, "DAEMONCORE  //  FIELDOPS", name="Aptos Display", size=7.5, bold=True, color=RED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(fp, "PRODUCTION QUICK REFERENCE  //  VERSION 6.4.0  //  AUTHORIZED OPERATIONS ONLY", name="Aptos", size=6.5, bold=True, color=MID)

# Page 1
label(doc, "Operator quick reference")
title = doc.add_paragraph(style="Title")
title.paragraph_format.space_after = Pt(2)
add_text(title, "DaemonCore FieldOps Production Cheat Sheet", name="Aptos Display", size=27, bold=True, color=BLACK)
title_props = title._p.get_or_add_pPr()
title_border = title_props.find(qn("w:pBdr"))
if title_border is not None:
    title_props.remove(title_border)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
add_text(subtitle, "Windows 6.4.0  |  Signed engagements  |  Evidence-first execution", name="Aptos Display", size=10, bold=True, color=RED)
intro(doc, "Use this sheet while operating FieldOps. The order matters: bind the operator, declare exact scope, verify the permit, execute only inside that scope, seal evidence, create findings, and close the engagement. A FieldOps license unlocks the workspace; it does not authorize a target.")

heading(doc, "The 60 second operating loop")
table(doc, ["Step", "Operator action", "Proof created"], [
    ("1  Bind", "Confirm the device-bound operator identity shows your legal or professional name, organization, email, and role.", "Ed25519 operator fingerprint"),
    ("2  Declare", "Create the engagement from the signed rules of engagement. Enter exact targets, ports, window, client, approver, and authorization reference.", "Signed operation permit"),
    ("3  Preflight", "Confirm ACTIVE AUTHORIZATION, target and port allowlists, policy level, execution profile, and permit validity.", "Scope and capacity record"),
    ("4  Execute", "Choose the workspace that matches the approved objective. Stop when the authorized proof point is reached.", "Operation receipt and output"),
    ("5  Seal", "Promote relevant results into the evidence vault. Import external JSON or SARIF only through Evidence Intake.", "SHA-256 capture digest"),
    ("6  Decide", "Create evidence-backed findings, record disposition, attach retests, and export the client report or case file.", "Finding history and report"),
    ("7  Close", "Stop active jobs, campaigns, and experiments, export the record, then close the authorization boundary.", "Closed immutable engagement"),
], [0.7, 4.35, 1.55])

heading(doc, "Hard stops")
table(doc, ["Stop immediately when", "Required response"], [
    ("The target, port, endpoint, method, or test window is not explicitly authorized", "Do not improvise. Pause and amend the rules of engagement."),
    ("DNS resolves to a private, loopback, link-local, reserved, or changed destination outside the declared mode", "Treat it as a scope change. Preserve the event and obtain authorization."),
    ("A campaign or workload crosses its approved objective, grant, or SLO", "Use the workspace stop control or emergency stop, then record the reason."),
    ("Evidence integrity or permit verification fails", "Do not rely on the record. Export what is available and escalate for review."),
], [3.15, 3.45])

# Page 2
page_break(doc)
label(doc, "Permit and preflight")
heading(doc, "Open a production engagement")
table(doc, ["Field", "Production entry", "Check"], [
    ("Engagement", "Specific job or assessment name", "No generic reusable permit"),
    ("Client", "Legal customer or internal owner", "Matches written authorization"),
    ("Authorization reference", "SOW, ROE, ticket, or change record", "Traceable by approver"),
    ("Network mode", "External or internal", "Matches the operator vantage point"),
    ("Targets", "Exact hostnames or IP addresses", "No inferred third parties"),
    ("TCP ports", "Exact approved port set", "Includes native tool and verification ports"),
    ("Policy", "Observe, Validate, or Stress", "Stress required for resilience authorization"),
    ("Execution profile", "Guarded or Professional", "Professional capacity follows signed scope"),
    ("Test window", "Start and end timestamps", "Ends within 366 days"),
    ("Approver", "Named authority and professional email", "Independent of operator where policy requires"),
], [1.35, 3.4, 1.8])

heading(doc, "Policy selection")
table(doc, ["Policy", "Permitted operation classes", "Use when"], [
    ("Observe", "Observation only", "Inventorying and recording approved signals without active validation"),
    ("Validate", "Observe and validate", "Confirming service, protocol, HTTP, TLS, and approved assessment evidence"),
    ("Stress", "Observe, validate, and resilience", "Running bounded Chaos Engine checks or verifying external load authority"),
], [1.0, 2.65, 2.9])

heading(doc, "Before every run")
bullet(doc, "Engagement status is active and the current clock is inside the signed test window.", "Window. ")
bullet(doc, "The selected target and every port are visible in the exact permit allowlist.", "Scope. ")
bullet(doc, "The operation class is allowed by the signed policy and the permit integrity indicator passes.", "Permit. ")
bullet(doc, "No conflicting native job, campaign, or Chaos Engine experiment is already active.", "Control plane. ")
bullet(doc, "The intended proof point, abort condition, communications path, and recovery owner are known.", "Exit plan. ")

# Page 3
page_break(doc)
label(doc, "War Room map")
heading(doc, "Choose the right workspace")
table(doc, ["Workspace", "Use it for", "Production output"], [
    ("Diagnostics", "DNS, TCP, HTTP, TLS, HTTP posture, service profiling, deep inventory, surface mapping, and baselines", "Sealed diagnostic capture"),
    ("Execution Fabric", "Managed Nmap inventory, workstation capability discovery, signed scope export, and external evidence intake", "Native job capture or signed manifest"),
    ("Campaigns", "Repeatable multi-target DNS, inventory, and surface workflows with pause, resume, cancel, and restart recovery", "Campaign task ledger and captures"),
    ("Assets and evidence", "Reviewing captures, service inventory, drift, imported JSON or SARIF, and promotion into findings", "Digest-verified evidence chain"),
    ("Findings", "Writing impact, remediation, severity, disposition, and retest history from selected evidence", "Client-ready finding record"),
    ("Chaos Engine", "Short bounded HEAD-based baseline, ramp, spike, and soak experiments with SLO abort and recovery proof", "Resilience score and experiment ledger"),
], [1.18, 3.6, 1.8])

heading(doc, "Managed Nmap run")
table(doc, ["Stage", "What to do"], [
    ("Select", "Choose one exact permitted target. FieldOps uses the engagement port set."),
    ("Attest", "Confirm the run remains inside the active signed permit."),
    ("Run", "Start Nmap and watch the live process output. FieldOps passes a fixed argument array without a command shell."),
    ("Stop", "Use STOP JOB when the objective changes, the window closes, or the target behaves unexpectedly."),
    ("Seal", "A completed result becomes a native-tool capture with engine, address, timing, output, and digest."),
], [1.1, 5.45])

heading(doc, "Evidence intake")
intro(doc, "Use signed scope export for Nuclei, ZAP, k6, Locust, TShark, Hashcat, and other supported bridges. Run the specialist tool in the customer-controlled environment, then import JSON or SARIF against the exact authorized target. Imported content is treated as evidence, never executed by DaemonCore.")

# Page 4
page_break(doc)
label(doc, "Verified Load Authority")
heading(doc, "High intensity testing without anonymous execution")
intro(doc, "Version 6.4.0 separates the built-in bounded Chaos Engine from professional external load generation. FieldOps signs k6 or Locust workload manifests only after the exact target publishes a challenge-bound capacity grant over trusted TLS. The customer-controlled runner performs the workload and returns results for evidence sealing.")

table(doc, ["Phase", "Required action", "Enforced boundary"], [
    ("1  Permit", "Create a new Stress-policy engagement containing the exact load target and verification port.", "Operator, approver, scope, and window are signed"),
    ("2  Publish", "Place the JSON capacity grant at the challenge path shown in Execution Fabric.", "Grant must originate from the exact target"),
    ("3  Verify", "Select target, port, and verified TLS, then choose Verify Target Capacity.", "Challenge, target, authorization reference, and expiration must match"),
    ("4  Configure", "Enter path, requests per second, duration, and concurrency at or below the grant.", "Values above the verified capacity are rejected"),
    ("5  Export", "Export the k6 or Locust signed workload manifest.", "Manifest includes permit and grant digests"),
    ("6  Execute", "Run through customer-controlled load infrastructure with monitoring and emergency stop available.", "Runner must independently enforce the manifest"),
    ("7  Import", "Return structured results through Evidence Intake and bind them to the target.", "Source digest enters chain of custody"),
], [1.0, 3.75, 1.85])

heading(doc, "Capacity grant schema")
table(doc, ["JSON field", "Meaning", "Rule"], [
    ("challenge", "Unique value shown by the engagement", "Exact match"),
    ("target", "Authorized hostname or IP", "Exact permit target"),
    ("authorizationReference", "SOW, ROE, ticket, or change record", "Exact engagement reference"),
    ("maxRequestsPerSecond", "Maximum approved arrival rate", "Positive integer"),
    ("maxDurationSeconds", "Maximum approved load duration", "At least 10 seconds"),
    ("maxConcurrency", "Maximum approved workers", "Positive integer"),
    ("validUntil", "Grant expiration in ISO 8601", "Inside the permit window"),
], [1.65, 3.1, 1.85])

intro(doc, "Emergency stop: stop the customer-controlled workers first, preserve target and runner telemetry, record the reason in the engagement, and do not resume until the approving authority confirms the remaining window and capacity.")

# Page 5
page_break(doc)
label(doc, "Evidence and incident card")
heading(doc, "Evidence to finding to report")
table(doc, ["Action", "Minimum production content"], [
    ("Capture", "Exact target, address, port or path, operation type, timestamps, duration, tool or engine, result, and digest"),
    ("Finding", "Specific title, severity, evidence-backed description, business impact, practical remediation, and source capture"),
    ("Disposition", "Open, accepted risk, resolved, or false positive, with operator note and timestamp"),
    ("Retest", "New capture from the same engagement, fixed or still-present verdict, note, and retained history"),
    ("Report", "Authorization, signed operator, approver, integrity state, campaigns, services, drift, evidence, findings, and retests"),
    ("Case file", "Portable engagement record retained before closure or upgrade"),
], [1.35, 5.2])

heading(doc, "Fast troubleshooting")
table(doc, ["Symptom", "Check first", "Next action"], [
    ("Operation blocked", "Entitlement, active window, policy class, exact target and port", "Correct the engagement; never bypass the permit"),
    ("Target resolves differently", "DNS answers and network mode", "Stop, record drift, and obtain a scope decision"),
    ("Nmap unavailable", "Execution Fabric workstation rescan and Docker availability", "Install the supported tool or use the signed evidence bridge"),
    ("Capacity verification fails", "New engagement challenge, exact JSON values, trusted TLS, HTTP 200, and grant expiration", "Correct the target-hosted grant and verify again"),
    ("Workload export blocked", "Stress policy, current grant, and requested rate, concurrency, and duration", "Stay at or below the target-issued capacity"),
    ("Evidence import rejected", "JSON or SARIF format, 2 MB limit, exact target, and source integrity", "Export structured evidence and retry"),
    ("Integrity indicator fails", "Permit mutation, file corruption, interrupted write, or stale record", "Stop execution and preserve the record for review"),
], [1.6, 2.65, 2.3])

heading(doc, "Closeout checklist")
table(doc, ["Done", "Closeout action"], [
    ("[ ]", "Stop all native jobs, campaigns, Chaos Engine experiments, and customer-controlled workers."),
    ("[ ]", "Confirm evidence, capture, and signature integrity indicators pass."),
    ("[ ]", "Finish finding dispositions and attach available retest evidence."),
    ("[ ]", "Export the client report and case file to the approved evidence location."),
    ("[ ]", "Record unresolved limitations, affected assets, owner, and next review date."),
    ("[ ]", "Close the authorization boundary. Reopen work only through a new engagement."),
], [.65, 5.9])

props = doc.core_properties
props.title = "DaemonCore FieldOps Production Cheat Sheet"
props.subject = "FieldOps 6.4.0 operator quick reference"
props.author = "DaemonCore Apps"
props.keywords = "DaemonCore, FieldOps, operator, production, cheat sheet"

doc.save(DOCX)
print(DOCX)
