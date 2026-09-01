from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "enterprise" / "DaemonCore-Academy-Enterprise-Training-Brief-v6.0.0.docx"
ASSETS = ROOT / "docs" / "marketing" / "fieldops-war-room"
ICON = ROOT / "build" / "icon-sizes" / "256.png"

RED = "E33E48"
DARK = "101216"
INK = "23272D"
MID = "616873"
LIGHT = "F2F4F6"
PALE_RED = "FBEDEF"
WHITE = "FFFFFF"
BORDER = "D9DDE2"
GREEN = "2F8F62"
CONTENT_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font(run, size=None, color=INK, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def hyperlink(paragraph, text, url, size, color=WHITE, bold=True, name="Aptos Display"):
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    node = OxmlElement("w:hyperlink")
    node.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    props.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    props.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    props.append(size_node)
    if bold:
        props.append(OxmlElement("w:b"))
    run.append(props)
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    node.append(run)
    paragraph._p.append(node)
    return node


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=100, start=130, bottom=100, end=130):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:tcMar"))
    if node is None:
        node = OxmlElement("w:tcMar")
        props.append(node)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = node.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            node.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    props = table._tbl.tblPr
    width_node = props.find(qn("w:tblW"))
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        props.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "130")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            value = widths[min(index, len(widths) - 1)]
            props = cell._tc.get_or_add_tcPr()
            cell_width = props.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                props.append(cell_width)
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = props.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        props.append(node)
    node.set(qn("w:val"), "true")


def configure_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)


def set_page_furniture(section, page):
    section.different_first_page_header_footer = False
    headers = (section.header, section.even_page_header, section.first_page_header)
    footers = (section.footer, section.even_page_footer, section.first_page_footer)
    for header in headers:
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font(paragraph.add_run("DAEMONCORE  //  ENTERPRISE TRAINING BRIEF"), 8, MID, True)
    for footer in footers:
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(paragraph.add_run(f"DAEMONCORE ACADEMY 6  |  SEPTEMBER 2026  |  {page}"), 8, MID, True)


def configure(doc):
    section = doc.sections[0]
    configure_geometry(section)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, RED),
        ("Heading 2", 13, 12, 6, RED),
        ("Heading 3", 12, 8, 4, DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    doc.core_properties.title = "DaemonCore Academy Enterprise Training & Workforce Readiness Brief"
    doc.core_properties.subject = "Enterprise cyber workforce development and authorized assessment capability"
    doc.core_properties.author = "DaemonCore Apps"
    doc.core_properties.keywords = "DaemonCore, enterprise training, cybersecurity workforce, cyber range, FieldOps"


def para(doc, text="", size=None, color=INK, bold=False, italic=False, align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    font(p.add_run(text), size, color, bold, italic)
    return p


def title_block(doc, number, kicker, title, lead):
    p = para(doc, f"{number}  //  {kicker.upper()}", 8.5, RED, True, after=5, keep=True)
    p.paragraph_format.keep_with_next = True
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    font(heading.add_run(title), 25, DARK, True, name="Aptos Display")
    para(doc, lead, 11.5, MID, after=13)


def image(doc, path, width, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt)
    return p


def caption(doc, text):
    return para(doc, text, 8, MID, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=9)


def metric_strip(doc, items, fill=DARK, value_size=18):
    table = doc.add_table(rows=1, cols=len(items))
    widths = [CONTENT_DXA // len(items)] * len(items)
    widths[-1] += CONTENT_DXA - sum(widths)
    table_geometry(table, widths)
    for cell, item in zip(table.rows[0].cells, items):
        value, label, *target = item
        shade(cell, fill)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        if target:
            hyperlink(cell.paragraphs[0], value, target[0], value_size)
        else:
            font(cell.paragraphs[0].add_run(value), value_size, WHITE, True, name="Aptos Display")
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(label.upper()), 7.5, "C7CBD1", True)
    return table


def callout(doc, label, text, fill=PALE_RED, accent=RED):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(label.upper()), 8, accent, True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(text), 10.3, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        font(p.add_run(item), 10.5, INK)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        font(p.add_run(item), 10.5, INK)


def data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_geometry(table, widths)
    repeat_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, DARK)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        font(cell.paragraphs[0].add_run(value.upper()), 8, WHITE, True)
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            if index % 2:
                shade(cell, LIGHT)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            font(cell.paragraphs[0].add_run(str(value)), 9.2, INK)
    table_geometry(table, widths)
    return table


def new_page(doc, page):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_geometry(section)
    set_page_furniture(section, page)


def build():
    doc = Document()
    configure(doc)

    # 1 - cover
    para(doc, "DAEMONCORE  //  ACADEMY", 10, RED, True, after=8)
    image(doc, ASSETS / "fieldops-war-room-hero.png", 6.5, "DaemonCore FieldOps War Room commercial campaign image")
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    font(title.add_run("ENTERPRISE TRAINING &\nWORKFORCE READINESS"), 29, DARK, True, name="Aptos Display")
    para(doc, "A practical cyber capability platform for teams that need more than passive course completion.", 13, MID, after=14)
    metric_strip(doc, [("127", "complete lessons"), ("120h45", "guided curriculum"), ("70", "specialist conditions"), ("7", "sealed field missions")])
    para(doc, "ACADEMY 6  //  ENTERPRISE CAPABILITY BRIEF  //  SEPTEMBER 2026", 8, MID, True, after=0)

    # 2 - executive brief
    new_page(doc, 2)
    title_block(doc, "01", "Executive brief", "Build capability that leaves evidence.", "DaemonCore Academy combines structured instruction, scored decisions, disposable live environments, durable operator records, and an authorization-bound professional assessment workspace in one local-first desktop application.")
    callout(doc, "The enterprise proposition", "Give employees a repeatable path from understanding a control, to practicing it, to proving an outcome - without turning training progress into an imaginary leaderboard or exposing practice targets to the public internet.")
    doc.add_heading("Designed for", level=2)
    bullets(doc, [
        "Security teams building common assessment, evidence, and reporting habits across mixed experience levels.",
        "Engineering organizations strengthening application, identity, cloud, Linux, container, detection, and software supply-chain capability.",
        "Leaders who need practical learning records and after-action evidence rather than attendance alone.",
        "Consulting and internal assurance teams that want training and professional authorized diagnostics in the same operating model.",
    ])
    doc.add_heading("What makes it different", level=2)
    para(doc, "The Academy teaches a method: establish scope, collect a defensible signal, test the boundary, preserve evidence, communicate impact, and verify remediation. FieldOps carries that method into real authorized engagements with signed permits, exact targets, sealed captures, findings, retests, and client-ready exports.")
    callout(doc, "Current commercial model", "The Academy is free to use. FieldOps is the paid professional gate. The current retail FieldOps license is $199 one time; organization-wide deployment, support, procurement, and commercial terms should be scoped separately.", fill=LIGHT, accent=DARK)

    # 3 - platform
    new_page(doc, 3)
    title_block(doc, "02", "Platform model", "One desktop. Four connected capability layers.", "Employees move through guided learning, live practice, scored proof, and professional operations without stitching together unrelated websites and disposable progress records.")
    image(doc, ASSETS / "field-kit-campaign.png", 6.15, "DaemonCore professional authorized assessment field kit")
    caption(doc, "A local-first operating model designed for dedicated workstations and controlled team environments.")
    data_table(doc, ["Layer", "Purpose", "Proof produced"], [
        ("Mission OS", "Diagnose strengths and select role-based routes", "Persisted baseline and next actions"),
        ("Academy + Forge", "Teach concepts through workshops and live specialist conditions", "Attempts, decisions, accepted conditions"),
        ("Sealed Range", "Practice unrestricted investigation inside disposable containment", "Evidence ledger and completion receipt"),
        ("FieldOps War Room", "Run authorized diagnostics and assessment campaigns", "Signed permits, captures, findings, reports"),
    ], [1700, 4050, 3610])

    # 4 - curriculum
    new_page(doc, 4)
    title_block(doc, "03", "Curriculum architecture", "Breadth for the workforce. Depth for specialists.", "The catalog is large enough to support foundational onboarding, role development, cross-training, and advanced evidence-led practice without forcing every employee through one identical route.")
    data_table(doc, ["Curriculum", "Lessons", "Time", "Focus"], [
        ("Full-Spectrum Security Assessment", "28", "24h", "Scope, networks, Windows, Linux, web, identity, cloud, containers, supply chain, evidence"),
        ("Web + API Specialist", "27", "24h45", "Browser contexts, authorization, OAuth, JWT, GraphQL, business logic, races, framing, caching"),
        ("Enterprise Forge pathways", "72", "72h", "AD, cloud engineering, detection, Linux, Kubernetes, software supply-chain defense"),
        ("Total", "127", "120h45", "Eight complete Academy pathways"),
    ], [2600, 900, 1050, 4810])
    doc.add_heading("Six Mission OS routes", level=2)
    bullets(doc, [
        "Penetration Tester - validate attack paths and write evidence-led findings.",
        "Web & API Specialist - identify and explain application trust failures.",
        "Identity Security - map authentication and delegated-control paths.",
        "Cloud Security - audit effective access and software provenance.",
        "Detection & Response - turn telemetry into bounded incident conclusions.",
        "Security Engineer - build controls across application, identity, and cloud systems.",
    ])
    callout(doc, "Advanced sequence", "Eight hour-long deep dives inside the full-spectrum path cover Linux privilege graphs, Windows service control, Active Directory paths, Kerberos trust, segmentation and pivot analysis, SSRF, OAuth/OIDC, and CI/CD provenance.")

    # 5 - workflow
    new_page(doc, 5)
    title_block(doc, "04", "Learning workflow", "Learn. Practice. Launch. Prove.", "Every activity identifies whether the operator is reading an example, making a scored decision, or working against a live disposable target. This removes the ambiguity common to terminal-styled courses.")
    numbered(doc, [
        "Diagnose - a twelve-scenario baseline measures six capability domains and persists to the local operator record.",
        "Select - Mission OS recommends a role route and calculates progress from completed work on the device.",
        "Learn - lessons combine mental models, commands, artifacts, workshops, validation checks, and primary references.",
        "Practice - Web Forge and Enterprise Forge require accepted live conditions rather than canned browser output.",
        "Launch - sealed Docker ranges open only after containment verification and provide an unrestricted in-range shell.",
        "Prove - after-action reviews score evidence coverage, independence, method discipline, and time discipline.",
    ])
    doc.add_heading("Recorded learner signals", level=2)
    metric_strip(doc, [("Attempts", "method repetition"), ("Evidence", "accepted outcomes"), ("Scores", "decision quality"), ("Time", "guided effort")], fill="2A2E34")
    para(doc, "The local operator record retains completion, attempts, practical scores, capstone decisions, weekly minutes, achievements, and activity. Progress is calculated from completed work; DaemonCore does not seed fake readiness scores or public leaderboards.", 10.5, MID, after=8)
    callout(doc, "Management implication", "Today, records are workstation-local and exportable. A centralized enterprise administration console, SSO/SCIM lifecycle, LMS synchronization, and organization-wide analytics are not represented as shipping capabilities.", fill=LIGHT, accent=DARK)

    # 6 - practice
    new_page(doc, 6)
    title_block(doc, "05", "Practical environments", "Practice against systems that can fail safely.", "DaemonCore separates high-intensity learning from public-network operations. The most permissive exercises stay inside disposable ranges with verified containment.")
    metric_strip(doc, [("48", "enterprise cases"), ("22", "web conditions"), ("7", "field missions"), ("3", "principal capstones")])
    doc.add_heading("Sealed range contract", level=2)
    bullets(doc, [
        "Internal-only Docker networking with no target ports published to the host.",
        "Zero host mounts, no privileged containers, dropped capabilities, no-new-privileges, and resource ceilings.",
        "Verified internet-egress denial before the operator shell is released.",
        "Run-specific seeds, full-tree pack fingerprints, launch receipts, and SHA-256 evidence ledgers.",
        "Guided, Assisted, Blind, and Professional modes with different hint availability and scoring behavior.",
    ])
    doc.add_heading("High-intensity resilience laboratory", level=2)
    para(doc, "Inside the sealed Academy range, bounded breakpoint experiments can run at up to 500 requests per second, 100 concurrent workers, and a 30,000-request budget against a disposable target that cannot reach the host or internet. The purpose is saturation, breakpoint, guardrail, and recovery engineering - not anonymous public-network disruption.")
    callout(doc, "Safety boundary", "DaemonCore does not provide an arbitrary public-network shell, DDoS engine, or online password-guessing system. Unrestricted command execution and high-intensity exercises remain inside the sealed range.")

    # 7 - FieldOps
    new_page(doc, 7)
    title_block(doc, "06", "FieldOps War Room", "Turn authorized assessments into reviewable operations.", "FieldOps is the professional layer for scoped diagnostics, multi-target campaigns, asset intelligence, sealed evidence, finding lifecycle, retesting, reporting, and bounded resilience work.")
    image(doc, ASSETS / "authorized-operations-team.png", 6.15, "Professional team using the DaemonCore FieldOps War Room")
    caption(doc, "Campaign imagery. The shipping War Room presents real operator, permit, target, evidence, finding, campaign, and ledger state.")
    data_table(doc, ["Capability", "Enterprise value"], [
        ("Live command deck", "Selected permit window, signed operator, integrity state, surface coverage, campaigns, findings, and latest ledger activity"),
        ("Campaign Engine", "Repeatable multi-target Complete Assessment, Service Inventory, and Change Verification profiles"),
        ("Evidence vault", "Digest-sealed raw captures with target, timing, resolved addresses, and immutable source context"),
        ("Findings and retests", "Severity, impact, remediation, disposition, and closure backed by later evidence"),
        ("Exports", "Machine-readable case file plus printable professional assessment report"),
    ], [2350, 7010])

    # 8 - governance
    new_page(doc, 8)
    title_block(doc, "07", "Authorization and governance", "Capability begins at a signed boundary.", "A FieldOps license unlocks the tool. It does not authorize a target. Execution requires an active engagement permit tied to a protected operator identity and exact scope.")
    doc.add_heading("Every new permit binds", level=2)
    bullets(doc, [
        "Named operator, organization, role, device-key fingerprint, and Ed25519 public key.",
        "Client or system owner, approving authority, approver email, and rules-of-engagement reference.",
        "Observe, Validate, or Stress policy level; internal or external network boundary.",
        "Exact targets, declared TCP ports, valid-from time, valid-until time, and operator attestation.",
    ])
    doc.add_heading("Execution safeguards", level=2)
    data_table(doc, ["Control", "Behavior"], [
        ("Destination pinning", "Resolve and pin an exact address before an operation begins"),
        ("Boundary rejection", "Block loopback, link-local, multicast, reserved, and mixed-boundary results"),
        ("Tool bridge", "Pass only normalized pinned addresses and declared ports through direct process invocation"),
        ("Evidence integrity", "Hash captures and chain audit events; preserve completed and blocked actions"),
        ("Attribution", "Sign new permits and operation receipts with the protected device key"),
    ], [2350, 7010])
    callout(doc, "Evidence claim", "Device-key signatures provide attribution and tamper evidence. They do not independently prove that a typed identity or authorization statement is truthful; organizations remain responsible for validating identity and authority.", fill=LIGHT, accent=DARK)

    # 9 - deployment
    new_page(doc, 9)
    title_block(doc, "08", "Deployment and procurement facts", "A clear current-state view for enterprise evaluation.", "DaemonCore is a local-first desktop platform. The following status is stated directly so technical, security, and procurement teams can evaluate the product without implied certifications or hidden cloud dependencies.")
    data_table(doc, ["Area", "Current state", "Enterprise consideration"], [
        ("Windows", "Stable 6.0 installer for Windows 10/11 x64", "Installer is not yet Authenticode-signed; verify release SHA-256 checksums"),
        ("Linux", "x64 AppImage and Debian package beta", "Ubuntu 22.04+ and Debian 12+ target; Linux remains separately labeled beta"),
        ("Runtime", "Docker Desktop on Windows; Docker Engine on Linux", "Dedicated workstation or managed VM recommended for live ranges"),
        ("Secrets", "OS-protected credentials; GNOME Keyring/KWallet on Linux", "FieldOps blocks unencrypted Electron basic_text storage"),
        ("Records", "Local operator and engagement data with exports", "Define endpoint backup, retention, and case-handling policy"),
        ("Administration", "No shipping central tenant administration console", "Plan workstation-based pilot; scope SSO, LMS, analytics, and fleet needs separately"),
        ("Assurance", "No SOC 2, ISO 27001, or FedRAMP claim is made", "Run normal vendor security and legal review before broad deployment"),
    ], [1500, 3200, 4660])
    callout(doc, "Recommended enterprise posture", "Begin with a managed pilot on dedicated endpoints or virtual machines, validate the release checksum, constrain Docker access to trained operators, and treat exported learner and engagement records according to company retention policy.")

    # 10 - pilot
    new_page(doc, 10)
    title_block(doc, "09", "Suggested enterprise pilot", "Prove the operating model before scaling it.", "A six-week pilot gives security, learning, endpoint, and procurement stakeholders enough evidence to judge instructional value, deployment fit, and governance requirements.")
    data_table(doc, ["Week", "Activity", "Exit evidence"], [
        ("0", "Security, legal, endpoint, and training readiness review", "Approved pilot boundary and endpoint standard"),
        ("1", "Install, checksum validation, Docker preflight, operator onboarding", "Working managed environment and named cohort"),
        ("2", "Mission OS diagnostic and route assignment", "Baseline across six capability domains"),
        ("3", "Role lessons, workshops, and specialist conditions", "Completion, attempt, and accepted-condition sample"),
        ("4", "Sealed range missions in Guided and Professional modes", "Evidence ledgers and after-action reviews"),
        ("5", "Principal capstone and remediation cycle", "Decision record, mastery gaps, reassigned work"),
        ("6", "Optional FieldOps tabletop on organization-owned test assets", "Signed permit, sealed captures, findings, and report"),
    ], [850, 4810, 3700])
    doc.add_heading("Pilot success measures", level=2)
    bullets(doc, [
        "Operators can distinguish worked examples from live actions without instructor intervention.",
        "At least one accepted evidence chain is produced in each assigned practical domain.",
        "Participants improve decision quality or reduce guidance use across repeat attempts.",
        "Endpoint, Docker, keyring, export, and upgrade procedures meet internal operational expectations.",
        "Security leadership can identify which centralized features are required before a larger rollout.",
    ])

    # 11 - roles
    new_page(doc, 11)
    title_block(doc, "10", "Stakeholder value", "One platform, different reasons to care.", "The enterprise decision is not only a training purchase. It touches capability development, endpoint operations, evidence handling, authorized testing, and procurement readiness.")
    data_table(doc, ["Stakeholder", "Primary value", "Decision evidence"], [
        ("CISO / security leadership", "Role-aligned capability development and defensible practice boundaries", "Pilot outcomes, gaps, and governance requirements"),
        ("Security managers", "Structured routes, repeat attempts, capstones, and after-action remediation", "Exported operator records and practical receipts"),
        ("Learning & development", "120h45m catalog with clear learning-to-practice workflow", "Cohort plan, completion signals, learner feedback"),
        ("Internal assurance / red team", "Signed scope, campaigns, evidence, findings, retests, reports", "FieldOps tabletop and sample case file"),
        ("Endpoint engineering", "Local deployment with controlled Docker and keyring dependencies", "Install, upgrade, rollback, and data-retention runbook"),
        ("Procurement / legal", "Transparent retail model and explicit current-state disclosures", "Commercial agreement, EULA, security review"),
    ], [1900, 4200, 3260])
    doc.add_heading("Where DaemonCore fits", level=2)
    para(doc, "DaemonCore is best positioned as a practical workforce-development and authorized assessment platform for teams willing to manage local endpoints and evaluate current enterprise gaps directly. It should complement - not falsely claim to replace - a company's identity provider, LMS, SIEM, vulnerability-management system, or centralized cyber range infrastructure.")
    callout(doc, "Buyer takeaway", "The near-term value is real practical depth, sealed local environments, evidence-led progression, and accountable FieldOps operations. The enterprise expansion opportunity is centralized administration, integrations, fleet deployment, and formal assurance.")

    # 12 - close
    new_page(doc, 12)
    para(doc, "DAEMONCORE  //  ENTERPRISE EVALUATION", 9, RED, True, after=10)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(9)
    font(heading.add_run("Build operators who can\nshow their work."), 30, DARK, True, name="Aptos Display")
    para(doc, "Use the platform to diagnose capability, build deliberate learning routes, practice inside disposable systems, preserve evidence, and graduate authorized operators into the FieldOps War Room.", 13, MID, after=18)
    callout(doc, "Recommended next step", "Run a managed six-week cohort with representative junior, mid-level, and senior practitioners. Include one endpoint engineer, one security leader, and one training owner in the evaluation team.")
    doc.add_heading("Evaluation checklist", level=2)
    bullets(doc, [
        "Select pilot roles, endpoints, and a named internal owner.",
        "Review EULA, release signing notice, Docker privileges, and local data handling.",
        "Choose Academy routes and define practical evidence expectations.",
        "Decide whether FieldOps is included and identify organization-owned test assets.",
        "Document required enterprise integrations and support terms before scale-up.",
    ])
    doc.add_paragraph()
    metric_strip(doc, [
        ("academy.daemoncore.app", "product and downloads", "https://academy.daemoncore.app"),
        ("github.com/DaemoncoreApps", "release transparency", "https://github.com/DaemoncoreApps"),
    ], fill=DARK, value_size=12.5)
    para(doc, "Product facts validated against the DaemonCore Academy 6.0.0 repository and release contract on 1 September 2026. This brief is product information, not a certification, penetration-test authorization, or substitute for organizational security review.", 8.2, MID, italic=True, after=0)

    # Word has no presentation-table role, so mark the first row of every data or
    # layout table to give assistive technology a stable point of reference.
    for table in doc.tables:
        repeat_header(table.rows[0])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", action="store_true")
    args = parser.parse_args()
    if not args.docx:
        parser.error("use --docx")
    build()
