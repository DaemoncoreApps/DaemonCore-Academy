from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "guides"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "DaemonCore-FieldOps-Operator-Fieldbook-v6.5.1.docx"

RED = "D71920"
BLACK = "101010"
DARK = "252525"
MID = "646464"
LIGHT = "F2F2F2"
PALE = "FAF6F6"
BORDER = "D9D9D9"
WHITE = "FFFFFF"


def set_font(run, name="Aptos", size=9, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def run(p, text, **kwargs):
    item = p.add_run(text)
    set_font(item, **kwargs)
    return item


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), color)


def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table):
    props = table._tbl.tblPr
    borders = props.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:color"), BORDER)
        borders.append(line)


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_borders(t)
    repeat_header(t.rows[0])
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, DARK)
        cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run(p, text.upper(), size=7.2, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run(p, str(value), size=7.7, bold=i == 0, color=BLACK if i == 0 else DARK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return t


def label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    run(p, text.upper(), name="Aptos Display", size=7.2, bold=True, color=RED)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(5)
    run(p, text, name="Aptos Display", size=18 if level == 1 else 11, bold=True)
    return p


def paragraph(doc, text, size=8.8, color=DARK, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead:
        run(p, bold_lead, size=size, bold=True, color=BLACK)
    run(p, text, size=size, color=color)
    return p


def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.19)
    p.paragraph_format.first_line_indent = Inches(-.12)
    p.paragraph_format.space_after = Pt(3)
    if lead:
        run(p, lead, size=8.2, bold=True)
    run(p, text, size=8.2, color=DARK)


def numbered(doc, text, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.05)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(3)
    if lead:
        run(p, lead, size=8.2, bold=True)
    run(p, text, size=8.2, color=DARK)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.14)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    run(p, text, name="Cascadia Mono", size=7.2, color=BLACK)


def new_page(doc, section, title):
    doc.add_page_break()
    label(doc, section)
    heading(doc, title)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(paragraph, "FIELDOPS // 6.5.1     ", size=6.7, bold=True, color=MID)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.append(begin)
    field_run._r.append(instr)
    field_run._r.append(separate)
    result_run = paragraph.add_run("1")
    set_font(result_run, size=6.7, bold=True, color=MID)
    end_run = paragraph.add_run()
    end_run._r.append(end)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(5.5)
section.page_height = Inches(8.5)
section.top_margin = Inches(.48)
section.bottom_margin = Inches(.48)
section.left_margin = Inches(.48)
section.right_margin = Inches(.48)
section.header_distance = Inches(.2)
section.footer_distance = Inches(.2)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
for style_name in ("Title", "Heading 1", "Heading 2"):
    styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
    styles[style_name].font.name = "Aptos Display"
    styles[style_name]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    styles[style_name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")

title_props = styles["Title"]._element.get_or_add_pPr()
title_border = title_props.find(qn("w:pBdr"))
if title_border is not None:
    title_props.remove(title_border)

footer = section.footer.paragraphs[0]
add_page_field(footer)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(12)
run(p, "DAEMONCORE", name="Aptos Display", size=10, bold=True, color=RED)
title = doc.add_paragraph(style="Title")
title.paragraph_format.space_after = Pt(8)
run(title, "FieldOps Operator Fieldbook", name="Aptos Display", size=28, bold=True)
paragraph(doc, "A practical desk and field reference for authorized assessment, evidence, resilience, and closeout work.", size=11, color=MID)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(26)
run(p, "RELEASE 6.5.1", size=8, bold=True, color=RED)
paragraph(doc, "Windows and Linux edition\nDaemonCore Apps\nSeptember 2026", size=9.2)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(44)
run(p, "CONTROLLED PROFESSIONAL USE", size=7.2, bold=True, color=RED)
paragraph(doc, "A license unlocks FieldOps. It does not authorize a target. Work only under valid written authorization, exact scope, an active testing window, and an agreed stop procedure.", size=8.4)

new_page(doc, "Operator Doctrine", "The rule that keeps every action defensible")
paragraph(doc, "FieldOps is most useful when the operator treats authorization, execution, evidence, and reporting as one continuous chain. The goal is not maximum activity. The goal is the smallest reliable proof that answers the approved question.")
for lead, text in [
    ("Know the proof point. ", "Write the condition that will end the test before execution starts."),
    ("Pin the destination. ", "Confirm the target, resolved address, port, protocol, and path against the signed permit."),
    ("Measure before pressure. ", "Capture a clean baseline and confirm target-side monitoring before a load profile."),
    ("Stop on uncertainty. ", "Treat DNS drift, unexpected ownership, telemetry loss, or a scope mismatch as a stop condition."),
    ("Seal as you go. ", "Promote useful captures while context is fresh; do not reconstruct the evidence chain later."),
]: bullet(doc, text, lead)
heading(doc, "Five questions before Run", 2)
add_table(doc, ["Question", "Required answer"], [
    ("Authority", "Who approved this exact action?"),
    ("Boundary", "Which target, address, port, path, and time window apply?"),
    ("Objective", "What observation proves or disproves the condition?"),
    ("Control", "Who can stop the activity and how will recovery be verified?"),
    ("Evidence", "What must be preserved for another reviewer to reproduce the conclusion?"),
], [1.05, 3.25])

new_page(doc, "Quick Start", "From zero to a defensible first capture")
for lead, text in [
    ("1  Verify", " the release checksum, launch on a dedicated operator account, and confirm version 6.5.1."),
    ("2  Activate", " FieldOps Pro and confirm secure credential storage is available."),
    ("3  Bind", " the named operator identity and review the displayed organization and email."),
    ("4  Declare", " the engagement from the written rules of engagement, using exact targets and ports."),
    ("5  Baseline", " with DNS, TCP, HTTP, TLS, or Complete Target Baseline before broader activity."),
    ("6  Seal", " material captures, create evidence-backed findings, and export the case file before closeout."),
]: numbered(doc, text, lead)
heading(doc, "Fast workstation checks", 2)
code(doc, "Windows:  Get-FileHash .\\DaemonCore-Academy-Setup.exe -Algorithm SHA256")
code(doc, "Linux:    sha256sum ./DaemonCore-Academy-*.AppImage")
code(doc, "Tools:    docker version    nmap --version    k6 version")
paragraph(doc, "Tip. Capture the tool version output in the case notes. It explains capability differences and makes later reproduction easier.", size=8.2, bold_lead="Operator note. ")

new_page(doc, "Preflight", "The sixty second go or no go")
add_table(doc, ["Check", "Go condition"], [
    ("Identity", "Correct named operator and protected device key"),
    ("Permit", "Signature valid; status active; policy permits the action"),
    ("Time", "Current time is inside the approved window"),
    ("Destination", "Exact target and port are listed; network mode matches"),
    ("Resolution", "Resolved addresses remain inside the declared boundary"),
    ("Dependencies", "Required Docker, Nmap, or k6 adapter reports ready"),
    ("Observers", "Owner, approver, and recovery contact know the start time"),
    ("Abort", "Local stop and target-side stop have named owners"),
    ("Evidence", "Clock, monitoring, log retention, and export location are ready"),
], [1.05, 3.25])
paragraph(doc, "No go means no improvisation. Correct the permit or environment, preserve the blocked event, and begin again from preflight.", bold_lead="Decision. ")

new_page(doc, "War Room", "Choose the workspace by the question")
add_table(doc, ["Workspace", "Use it when you need"], [
    ("Diagnostics", "One focused DNS, TCP, HTTP, TLS, service, inventory, or surface answer"),
    ("Execution Fabric", "Native tool readiness, managed Nmap, verified k6, scope export, or evidence intake"),
    ("Campaigns", "Repeatable approved checks across several declared targets"),
    ("Assets and evidence", "Service history, change intelligence, capture integrity, and promotion"),
    ("Findings", "Impact, severity, remediation, disposition, and retest history"),
    ("Chaos Engine", "A short bounded resilience sample with automatic SLO aborts"),
], [1.35, 2.95])
heading(doc, "Navigation trick", 2)
paragraph(doc, "Begin in the Engagement Vault, confirm the scope strip, then choose the workspace. Return to Assets and Evidence after every meaningful operation. That rhythm prevents unreviewed output from piling up.")

new_page(doc, "Diagnostics", "Select the least invasive useful operation")
add_table(doc, ["Need", "Start with", "Escalate to"], [
    ("Name and address truth", "DNS resolution", "DNS record profile"),
    ("One service reachable", "TCP reachability", "Service profile"),
    ("Web control posture", "HTTP response", "HTTP security posture"),
    ("Certificate identity", "TLS identity", "Service profile"),
    ("Declared ports", "Allowed port survey", "Deep service inventory"),
    ("Change over time", "Complete target baseline", "Repeat baseline and review drift"),
    ("Known control paths", "Bounded web map", "External approved specialist evidence"),
], [.95, 1.42, 1.93])
paragraph(doc, "A service fingerprint is a lead, not a finding. Confirm the observed behavior, environmental context, prerequisites, and business effect before assigning severity.", bold_lead="Evidence rule. ")

new_page(doc, "Engagement Design", "Make the permit operational")
for lead, text in [
    ("Targets. ", "Use exact hostnames or IP addresses. Separate internal and external boundaries."),
    ("Ports. ", "List only the ports authorized for this engagement; do not use a remembered standard list."),
    ("Window. ", "Include timezone, maintenance exclusions, and the latest permitted stop time."),
    ("Approver. ", "Record a named authority, professional email, and authorization reference that another reviewer can retrieve."),
    ("Policy. ", "Use Observe for focused collection, Validate for broader assessment, and Stress only for explicitly approved resilience work."),
    ("Capacity. ", "Professional mode signs the declared target and port lists as capacity; target-hosted grants separately govern managed load."),
]: bullet(doc, text, lead)
heading(doc, "Amend instead of stretching", 2)
paragraph(doc, "A new hostname, port, path, technique, date, or approving authority is a scope change. Close or supersede the permit and issue a new one. Never reinterpret an old permit during execution.")

new_page(doc, "Execution Fabric", "Use native tools without losing the chain")
for lead, text in [
    ("Rescan capabilities. ", "Do this after installing Docker, Nmap, or k6, and after changing PATH."),
    ("Prefer managed execution. ", "When available, let FieldOps pin arguments, record versions, stream output, and seal the result."),
    ("Use signed scope export. ", "For a specialist tool, export the declared boundary and preserve the manifest with the tool output."),
    ("Import structured evidence. ", "Use supported JSON or SARIF and bind it to the exact engagement target."),
    ("Preserve originals. ", "Keep the raw source artifact beside the FieldOps case export and record its SHA-256 digest."),
]: bullet(doc, text, lead)
heading(doc, "Dependency triage", 2)
add_table(doc, ["State", "First action"], [
    ("Docker unavailable", "Start Docker Engine or Desktop; verify daemon access"),
    ("Nmap unavailable", "Install Nmap or use the supported Docker adapter"),
    ("k6 unavailable", "Install Grafana k6; restart FieldOps; rescan capabilities"),
    ("Tool starts then stops", "Review timeout, permit window, target pin, and local process output"),
], [1.35, 2.95])

new_page(doc, "Verified Load", "Managed k6 launch sequence")
for lead, text in [
    ("1  Baseline. ", "Capture steady-state latency, error rate, saturation, and recovery behavior."),
    ("2  Publish. ", "Have the system owner place the challenge-bound capacity grant on the exact trusted TLS target."),
    ("3  Verify. ", "Confirm target, authorization reference, challenge, expiration, and signature-bound grant fields."),
    ("4  Configure. ", "Choose the approved baseline, ramp, spike, soak, breakpoint, or recovery profile within the grant."),
    ("5  Observe. ", "Keep application, infrastructure, dependency, queue, and user-impact telemetry visible."),
    ("6  Execute. ", "Arm managed k6 only after both local and target-side stop owners are ready."),
    ("7  Recover. ", "Continue observation after traffic ends until agreed recovery indicators return to baseline."),
    ("8  Seal. ", "Preserve the permit, grant, workload plan, engine version, telemetry summary, stop reason, and result digest."),
]: numbered(doc, text, lead)
paragraph(doc, "The verified grant is the capacity contract. If the intended rate, concurrency, duration, or expiration exceeds it, obtain a new grant rather than splitting the run.", bold_lead="Hard rule. ")

new_page(doc, "Live Control", "Read the system while pressure is active")
add_table(doc, ["Signal", "Interpretation", "Action"], [
    ("Achieved RPS", "Delivery versus requested rate", "Check dropped iterations and client limits"),
    ("P95 latency", "Tail response under load", "Compare with SLO and baseline"),
    ("Error rate", "Failed user-facing work", "Stop at approved abort threshold"),
    ("Saturation", "CPU, memory, pools, queues", "Find the first constrained dependency"),
    ("Dropped iterations", "Generator could not maintain schedule", "Separate generator limits from target limits"),
    ("Recovery time", "Return to steady state", "Do not declare success at traffic stop"),
], [1.1, 1.65, 1.55])
heading(doc, "Use the emergency stop when", 2)
bullet(doc, "The permit, grant, target identity, monitoring view, or communications path becomes uncertain.")
bullet(doc, "An agreed SLO, infrastructure threshold, user-impact threshold, or business stop condition is reached.")
bullet(doc, "The workload deviates from the signed plan or the target-side owner calls stop.")

new_page(doc, "Evidence", "Build a record another operator can trust")
add_table(doc, ["Preserve", "Why it matters"], [
    ("Permit and grant", "Proves the declared authority and capacity context"),
    ("Target and resolution", "Shows where execution was actually pinned"),
    ("Tool and version", "Supports repeatability and explains parser behavior"),
    ("Raw capture", "Retains the original observation before interpretation"),
    ("Timestamps and timezone", "Aligns endpoint, target, and monitoring events"),
    ("Digest and receipt", "Detects later change and binds the operation record"),
    ("Operator notes", "Separates observation, inference, decision, and limitation"),
], [1.45, 2.85])
paragraph(doc, "Write observations in present tense and inferences as conclusions supported by named evidence. If evidence is incomplete, state the limitation instead of filling the gap.")
heading(doc, "Minimal note pattern", 2)
code(doc, "OBSERVED: <fact>  SOURCE: <capture>  AT: <time>\nINFERRED: <meaning>  LIMIT: <uncertainty>  NEXT: <retest>")

new_page(doc, "Findings", "Turn captures into decisions")
for lead, text in [
    ("Title. ", "Name the failed control or exposed condition, not the scanner signature."),
    ("Evidence. ", "Attach the smallest capture set that proves the condition and preserves context."),
    ("Impact. ", "Describe the credible business or security consequence under observed prerequisites."),
    ("Severity. ", "Account for exposure, asset role, data sensitivity, controls, detection, and recovery."),
    ("Remediation. ", "Recommend a control outcome, owner, and validation step rather than a generic patch phrase."),
    ("Retest. ", "Repeat the minimum proof, confirm intended behavior still works, and attach a new sealed capture."),
]: bullet(doc, text, lead)
heading(doc, "Status discipline", 2)
add_table(doc, ["State", "Use when"], [
    ("Open", "The condition is validated and awaits action"),
    ("Accepted", "Authorized risk owner accepts the documented residual risk"),
    ("Remediated", "A change is reported but still needs independent retest"),
    ("Closed", "Retest evidence supports closure"),
], [1.05, 3.25])

new_page(doc, "Campaigns", "Scale repetition without losing control")
for lead, text in [
    ("Use one objective. ", "A campaign should answer one repeatable question across the selected targets."),
    ("Order low to high. ", "Resolve and baseline first; inventory or deeper validation follows only when needed."),
    ("Watch task state. ", "Pause on systematic failure instead of allowing bad assumptions to multiply."),
    ("Resume deliberately. ", "After restart recovery, confirm the permit window and target state before resuming pending tasks."),
    ("Review outliers. ", "A single different result may be the most important campaign outcome."),
]: bullet(doc, text, lead)
heading(doc, "Campaign stop conditions", 2)
paragraph(doc, "Repeated DNS drift, widespread timeouts, unexpected rate limiting, authentication prompts, owner stop request, or evidence-integrity failure should pause the campaign for human review.")

new_page(doc, "Troubleshooting", "Recover the workflow without guessing")
add_table(doc, ["Symptom", "Check first", "Next move"], [
    ("Operation blocked", "Permit status, policy, target, port, time", "Correct scope; preserve blocked event"),
    ("Wrong boundary", "DNS answers and Internal or External mode", "Stop; obtain a scope decision"),
    ("No native adapter", "PATH, process restart, capability rescan", "Install supported tool or use evidence bridge"),
    ("Capacity rejected", "Challenge, TLS, target, grant expiry and fields", "Republish the exact grant"),
    ("Load will not arm", "Stress policy, verified grant, k6 availability", "Resolve the failed preflight item"),
    ("Evidence rejected", "Format, size, exact target, source digest", "Re-export structured JSON or SARIF"),
    ("Integrity failed", "Original artifact and last known export", "Stop relying on record; escalate review"),
    ("License offline", "Last validation and grace deadline", "Restore connectivity and validate"),
], [1.18, 1.55, 1.57])

new_page(doc, "Stop and Escalate", "When professional judgment overrides momentum")
for text in [
    "A hostname resolves outside the approved boundary or changes ownership during the window.",
    "The target-side owner cannot see the activity, cannot stop it, or loses monitoring.",
    "Production health degrades beyond the agreed threshold or dependent services show collateral impact.",
    "The app reports permit, signature, grant, evidence, or secure-storage integrity failure.",
    "A technique requires credentials, destructive changes, persistence, data access, or exploitation not named in the rules of engagement.",
    "The operator cannot distinguish a target failure from a generator, network, or observation failure.",
]: bullet(doc, text)
heading(doc, "Stop record", 2)
code(doc, "TIME / OPERATOR / RUN ID / STOP SOURCE / OBSERVED CONDITION\nLOCAL ACTION / TARGET ACTION / RECOVERY OWNER / RECOVERY PROOF")

new_page(doc, "Closeout", "Leave the client with a usable record")
for lead, text in [
    ("Freeze execution. ", "Stop active jobs, confirm recovery, and close the authorization boundary when work ends."),
    ("Verify integrity. ", "Review capture digests, signed receipts, audit state, and imported evidence provenance."),
    ("Reconcile findings. ", "Every material observation is promoted, dismissed with rationale, or recorded as a limitation."),
    ("Export both views. ", "Preserve the machine-readable case file and the human-readable report."),
    ("Transfer securely. ", "Use the client's approved channel and document receipt, retention, and deletion expectations."),
    ("Back up before upgrade. ", "Keep an export outside the application data directory before changing workstation state."),
]: bullet(doc, text, lead)
heading(doc, "Final review", 2)
paragraph(doc, "Can a reviewer identify who acted, under whose authority, against which exact destination, with which tool and settings, what was observed, why the conclusion follows, and how recovery or remediation was verified? If not, the case is not finished.")

new_page(doc, "Pocket Reference", "The operator card")
label(doc, "Before")
paragraph(doc, "Identity - permit - time - target - port - resolution - policy - objective - observers - abort - evidence")
label(doc, "During")
paragraph(doc, "Watch destination - output - RPS - p95 - errors - saturation - dropped work - communications - stop conditions")
label(doc, "After")
paragraph(doc, "Recovery - capture - digest - receipt - finding - retest - case export - report - transfer - scope closure")
heading(doc, "Field shortcuts", 2)
add_table(doc, ["If you need", "Do this"], [
    ("A quick answer", "Run one focused diagnostic"),
    ("A repeatable picture", "Establish and repeat a complete baseline"),
    ("Many approved targets", "Use a campaign with one objective"),
    ("Specialist tooling", "Export signed scope and import structured evidence"),
    ("Real load evidence", "Verify target capacity and use managed k6"),
    ("A client decision", "Promote evidence into a finding and retest"),
], [1.4, 2.9])
paragraph(doc, "Authorized. Attributable. Reproducible.", size=10, bold_lead="The standard. ")

doc.core_properties.title = "DaemonCore FieldOps Operator Fieldbook"
doc.core_properties.subject = "Field reference for authorized FieldOps operations"
doc.core_properties.author = "DaemonCore Apps"
doc.core_properties.keywords = "DaemonCore, FieldOps, operator, evidence, resilience, k6"
doc.core_properties.comments = "Production release 6.5.1"
doc.save(DOCX)
print(DOCX)
