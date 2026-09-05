from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "guides" / "DaemonCore-FieldOps-Real-World-Casebook-v6.5.1.docx"
OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

RED = "D71920"
BLACK = "111111"
DARK = "292929"
MID = "666666"
LIGHT = "F2F2F2"
BORDER = "D9D9D9"
WHITE = "FFFFFF"


SCENARIOS = [
    {
        "code": "CASE 01", "title": "Shadow Port", "subtitle": "External attack surface drift after an emergency deployment",
        "client": "Northstar Benefits", "role": "External assessor", "difficulty": "Foundation", "time": "60 to 90 minutes", "policy": "Validate",
        "brief": "Northstar's infrastructure team completed an emergency customer-portal deployment overnight. A monitoring analyst now sees a service on a port that is not present in the approved architecture record. The release team says the exposure is temporary; the system owner wants a defensible answer before the morning change meeting.",
        "pressure": "The change manager wants a verbal answer in twenty minutes and suggests scanning adjacent ports to save time.",
        "scope": [("Targets", "Facilitator-assigned portal hostname and pinned lab address"), ("Ports", "443 and one facilitator-assigned candidate port"), ("Window", "Current exercise window only"), ("Excluded", "Adjacent hosts, credentials, exploitation, configuration changes"), ("Stop", "Ownership mismatch, DNS drift, monitoring loss, or client stop")],
        "objectives": ["Confirm the authorization boundary before any diagnostic.", "Establish a current surface baseline and compare it with available prior evidence.", "Determine what can be stated about the unexpected service without overstating the evidence.", "Create a finding or documented non-finding with a reproducible evidence trail.", "Define a minimal remediation retest."],
        "artifacts": ["Approved target and port list", "Prior architecture extract", "Overnight change ticket summary", "Current DNS response", "Blank client finding template"],
        "injects": [("T plus 10", "DNS returns an additional address not mentioned in the briefing.", "Decide whether execution may continue and record why."), ("T plus 25", "The release engineer says the port belongs to a temporary admin service.", "Separate testimony from observed evidence."), ("T plus 40", "The service stops responding before the review is complete.", "Preserve the transient evidence and design the retest.")],
        "evidence": ["Permit status and scope strip", "DNS and resolution record", "Allowed-port or complete surface capture", "Service identity evidence", "Prior-versus-current comparison", "Decision and limitation note"],
        "deliverable": "One-page exposure decision, supporting capture list, finding disposition, and remediation retest plan.",
        "rubric": [("Boundary control", 20), ("Diagnostic selection", 15), ("Evidence integrity", 20), ("Reasoning and limitations", 20), ("Finding quality", 15), ("Retest plan", 10)],
    },
    {
        "code": "CASE 02", "title": "Launch Day", "subtitle": "API degradation during a planned traffic event",
        "client": "Atlas Ticketing", "role": "Resilience operator", "difficulty": "Intermediate", "time": "90 to 120 minutes", "policy": "Stress",
        "brief": "Atlas expects a large on-sale event tomorrow. The application team increased API capacity but has never measured the new breakpoint or recovery time. The owner has approved a controlled test against a dedicated production-like endpoint and will publish the target-hosted capacity grant.",
        "pressure": "Leadership asks for the highest possible request rate and wants the run repeated immediately if the first result is unclear.",
        "scope": [("Target", "Facilitator-assigned HTTPS API endpoint"), ("Path", "Assigned synthetic health or catalog path"), ("Method", "Managed k6 HTTP GET profile only"), ("Capacity", "Exactly the verified target-hosted grant"), ("Excluded", "Authentication, purchase submission, destructive data, ungranted workers")],
        "objectives": ["Capture a clean baseline before pressure.", "Verify the challenge-bound capacity grant and select an appropriate profile.", "Correlate achieved RPS, p95 latency, errors, dropped iterations, and target saturation.", "Apply the agreed stop criteria without waiting for total failure.", "Prove recovery and produce a capacity conclusion that distinguishes target limits from generator limits."],
        "artifacts": ["Rules of engagement", "Capacity grant JSON", "Service SLO", "Target monitoring snapshot", "Change calendar", "Recovery-owner contact card"],
        "injects": [("T plus 15", "The grant expires sooner than the planned workload would finish.", "Resolve the mismatch before arming."), ("T plus 35", "P95 rises while error rate remains low and a queue approaches saturation.", "Decide whether the abort threshold or proof point has been reached."), ("T plus 55", "Traffic ends but queue depth does not return to baseline.", "Continue recovery observation and avoid premature success.")],
        "evidence": ["Baseline capture", "Verified grant digest", "Signed workload plan", "k6 engine version", "Live telemetry summary", "Stop reason", "Recovery proof"],
        "deliverable": "Capacity and recovery report stating the tested envelope, first degradation signal, stop decision, recovery time, and next safe experiment.",
        "rubric": [("Grant verification", 20), ("Profile design", 15), ("Telemetry interpretation", 25), ("Stop discipline", 15), ("Recovery proof", 15), ("Report precision", 10)],
    },
    {
        "code": "CASE 03", "title": "Broken Chain", "subtitle": "TLS identity failure after certificate renewal",
        "client": "Meridian Health Logistics", "role": "Application security assessor", "difficulty": "Foundation", "time": "60 to 90 minutes", "policy": "Observe",
        "brief": "A partner integration began rejecting one regional endpoint after a certificate renewal. Browser users report mixed results. Operations believes the certificate is valid because the primary endpoint works. The client needs to know whether the failure is identity, chain, protocol, load-balancer, or client-specific behavior.",
        "pressure": "A manager asks the operator to ignore hostname validation temporarily so the partner can reconnect.",
        "scope": [("Targets", "Two facilitator-assigned regional hostnames"), ("Ports", "443 only"), ("Permitted", "DNS, TCP, TLS identity, and HTTP response evidence"), ("Excluded", "Certificate changes, authentication, private keys, bypass advice"), ("Stop", "Unexpected ownership or resolution outside assigned boundary")],
        "objectives": ["Compare name resolution and pinned destinations for both endpoints.", "Capture certificate identity, protocol, and chain-relevant evidence.", "Distinguish observed facts from client compatibility hypotheses.", "Identify the minimum corrective control and a safe validation sequence.", "Document residual uncertainty where FieldOps evidence is insufficient."],
        "artifacts": ["Partner error excerpt", "Renewal change record", "Expected hostnames", "Regional load-balancer inventory", "Approved maintenance window"],
        "injects": [("T plus 15", "One hostname alternates between two addresses.", "Capture resolution and avoid assuming equivalent configuration."), ("T plus 30", "The certificate dates are valid but the identity differs on one address.", "Describe the failed control precisely."), ("T plus 45", "Operations proposes disabling validation at the client.", "Recommend a control-preserving response.")],
        "evidence": ["DNS records", "TCP reachability", "TLS identity per pinned destination", "HTTP response metadata", "Comparison note", "Retest checklist"],
        "deliverable": "TLS incident brief with affected path, observed identity, likely control failure, immediate containment recommendation, and retest criteria.",
        "rubric": [("Comparative method", 20), ("TLS evidence", 20), ("Root-cause reasoning", 20), ("Control-preserving advice", 20), ("Retest quality", 10), ("Communication", 10)],
    },
    {
        "code": "CASE 04", "title": "Branchline", "subtitle": "Multi-site network control consistency review",
        "client": "Redwood Manufacturing", "role": "Network security consultant", "difficulty": "Intermediate", "time": "90 to 120 minutes", "policy": "Validate",
        "brief": "Redwood operates four plants that should share a common remote-access and management baseline. An audit found inconsistent firewall documentation. The client wants evidence of externally observable consistency before a wider remediation program begins.",
        "pressure": "The network lead says a timeout should be counted as secure and asks for a green summary by site.",
        "scope": [("Targets", "Four facilitator-assigned site endpoints"), ("Ports", "Exact facilitator-assigned management and web ports"), ("Permitted", "Campaign DNS, inventory, service, TLS, and posture checks"), ("Excluded", "Credential testing, exploitation, unlisted ranges"), ("Stop", "Systematic timeout, DNS drift, rate limiting, or site-owner request")],
        "objectives": ["Create a campaign that answers one consistent control question.", "Differentiate closed, filtered, unreachable, and unobserved states.", "Identify material outliers without treating every difference as a vulnerability.", "Preserve site-specific evidence and campaign-level reasoning.", "Prioritize remediation by control risk and business context."],
        "artifacts": ["Expected site standard", "Plant criticality table", "Approved targets and ports", "Network-owner roster", "Known maintenance exceptions"],
        "injects": [("T plus 20", "One site returns widespread timeouts after several completed tasks.", "Pause or continue based on evidence and communications."), ("T plus 40", "A site owner identifies an approved exception absent from the initial brief.", "Treat the exception as new context, not retroactive proof."), ("T plus 60", "Two services share a version string but expose different posture.", "Avoid version-only conclusions.")],
        "evidence": ["Campaign definition", "Per-target task ledger", "Service and posture captures", "Exception record", "Outlier matrix", "Campaign stop or resume decision"],
        "deliverable": "Four-site consistency matrix, prioritized exceptions, evidence map, and recommended validation sequence.",
        "rubric": [("Campaign design", 20), ("State interpretation", 20), ("Outlier analysis", 20), ("Exception handling", 15), ("Evidence mapping", 15), ("Prioritization", 10)],
    },
    {
        "code": "CASE 05", "title": "New Blood", "subtitle": "Post-acquisition perimeter review with incomplete records",
        "client": "Kestrel Financial Group", "role": "Integration security assessor", "difficulty": "Advanced", "time": "120 minutes", "policy": "Validate",
        "brief": "Kestrel acquired a small software company. The seller provided partial infrastructure records, and several services are operated by third parties. The integration team needs a defensible starting inventory without testing assets that the acquired company does not control.",
        "pressure": "The deal team provides a broad brand domain and says anything related to the name is probably in scope.",
        "scope": [("Targets", "Only facilitator-confirmed exact hostnames"), ("Ports", "Only listed service ports"), ("Permitted", "Ownership validation, surface baseline, service inventory, posture"), ("Excluded", "Related domains, third-party SaaS, discovered hosts, credentials"), ("Stop", "Ownership ambiguity, third-party indicators, or boundary mismatch")],
        "objectives": ["Convert incomplete records into exact testable scope without inference-based expansion.", "Establish an initial asset and service baseline.", "Track ownership confidence separately from technical observation.", "Identify evidence gaps that block a security conclusion.", "Produce an integration backlog ordered by ownership and exposure risk."],
        "artifacts": ["Partial asset spreadsheet", "Domain registration excerpt", "Vendor list", "Acquisition contact tree", "Approved exact-target addendum"],
        "injects": [("T plus 20", "A discovered hostname appears branded but resolves to a major SaaS provider.", "Exclude it until authority is confirmed."), ("T plus 45", "The seller's engineer claims ownership but cannot provide a contract or account record.", "Record confidence and escalate authority verification."), ("T plus 70", "A known asset exposes a service absent from all records.", "Validate only within declared ports and record the inventory gap.")],
        "evidence": ["Authorization addendum", "Ownership-confidence ledger", "Surface baselines", "Service inventory", "Third-party exclusions", "Evidence-gap register"],
        "deliverable": "Acquisition perimeter baseline, ownership-confidence matrix, excluded-asset log, and prioritized integration backlog.",
        "rubric": [("Scope discipline", 25), ("Ownership reasoning", 20), ("Inventory quality", 20), ("Gap management", 15), ("Risk prioritization", 10), ("Record clarity", 10)],
    },
    {
        "code": "CASE 06", "title": "Redline", "subtitle": "Production outage triage under conflicting signals",
        "client": "Harbor Freight Systems", "role": "Incident support operator", "difficulty": "Advanced", "time": "90 minutes", "policy": "Observe",
        "brief": "Customers intermittently fail to reach a shipping-status application. Internal dashboards show healthy application instances, while support reports regional failures. FieldOps is authorized for focused external diagnostics only; the incident commander owns all production changes.",
        "pressure": "Several teams offer competing explanations and ask the operator to prove one quickly.",
        "scope": [("Target", "Facilitator-assigned public application hostname"), ("Ports", "443 only"), ("Permitted", "DNS, TCP, TLS, HTTP, and repeated bounded baseline evidence"), ("Excluded", "Load generation, configuration changes, internal systems"), ("Stop", "Incident commander request, ownership drift, or material user impact")],
        "objectives": ["Build a time-aligned external observation sequence.", "Distinguish DNS, reachability, TLS, HTTP, dependency, and saturation hypotheses.", "Communicate what FieldOps can and cannot prove from the assigned vantage point.", "Preserve observations without delaying incident recovery.", "Recommend the next highest-information check to the incident commander."],
        "artifacts": ["Incident timeline excerpt", "Regional support reports", "Application dashboard snapshot", "Recent DNS change", "Incident command roles"],
        "injects": [("T plus 10", "DNS answers change while the application dashboard stays green.", "Align external evidence with the change timeline."), ("T plus 30", "TLS succeeds but HTTP responses alternate between success and gateway error.", "Avoid collapsing layers into one cause."), ("T plus 50", "The incident commander restores service before root cause is proven.", "Prioritize recovery evidence and preserve uncertainty.")],
        "evidence": ["Timestamped resolution series", "Layered diagnostic captures", "Observed-versus-reported matrix", "Decision log", "Recovery confirmation", "Unresolved hypotheses"],
        "deliverable": "Concise incident support timeline, layer-by-layer findings, confidence statement, and recommended follow-up evidence.",
        "rubric": [("Layered reasoning", 25), ("Time alignment", 20), ("Incident discipline", 20), ("Confidence language", 15), ("Recovery evidence", 10), ("Next action", 10)],
    },
    {
        "code": "CASE 07", "title": "Silent Alarm", "subtitle": "Blue-team visibility validation during controlled activity",
        "client": "Crown Peak Energy", "role": "Purple-team operator", "difficulty": "Intermediate", "time": "90 to 120 minutes", "policy": "Validate",
        "brief": "Crown Peak recently deployed new network detections. The defensive team wants to validate visibility of approved discovery and service-assessment activity without turning the exercise into an adversary emulation campaign.",
        "pressure": "The detection engineer asks for more aggressive behavior after the first activity produces no alert.",
        "scope": [("Targets", "Two facilitator-assigned monitored endpoints"), ("Ports", "Exact assigned service ports"), ("Permitted", "Named diagnostics, managed inventory, one approved campaign"), ("Excluded", "Evasion, exploitation, credentials, persistence, unapproved techniques"), ("Stop", "Monitoring loss, target health concern, or exercise-control request")],
        "objectives": ["Agree on expected telemetry before execution.", "Run a reproducible sequence that defenders can correlate.", "Separate missing telemetry, missing analytic logic, and missing alert routing.", "Change one variable at a time when repeating activity.", "Produce an evidence-backed detection improvement record."],
        "artifacts": ["Detection objective", "Expected data-source list", "Exercise communication plan", "Approved operation sequence", "Defender observation worksheet"],
        "injects": [("T plus 15", "Endpoint telemetry appears but the network sensor shows nothing.", "Determine whether execution or collection changed."), ("T plus 35", "An alert fires ten minutes late with the wrong asset owner.", "Record detection, enrichment, and routing separately."), ("T plus 55", "A defender requests an unapproved technique to force an alert.", "Hold the boundary and propose an amendment.")],
        "evidence": ["Signed permit", "Operation receipts", "Exact execution timeline", "Defender event identifiers", "Expected-versus-observed matrix", "Repeat-test plan"],
        "deliverable": "Detection validation record covering data collection, analytic result, enrichment, routing, timing, and recommended retest.",
        "rubric": [("Expectation design", 15), ("Reproducibility", 20), ("Correlation quality", 20), ("Gap classification", 20), ("Scope control", 15), ("Retest plan", 10)],
    },
    {
        "code": "CASE 08", "title": "Unknown Edge", "subtitle": "Emergency investigation of a reported internet exposure",
        "client": "Orchid Public Services", "role": "On-call security assessor", "difficulty": "Advanced", "time": "75 to 90 minutes", "policy": "Validate",
        "brief": "A researcher reports an administrative interface associated with Orchid. The screenshot lacks a reliable timestamp and the hostname is not in the asset inventory. Legal has authorized validation of one exact hostname after the infrastructure owner confirms control.",
        "pressure": "Executives want proof before notifying the vendor, while the reporter asks whether Orchid has started testing.",
        "scope": [("Target", "Single facilitator-assigned exact hostname"), ("Ports", "443 and one assigned administrative port"), ("Permitted", "Ownership, DNS, reachability, TLS, HTTP posture, bounded path map"), ("Excluded", "Login attempts, credentials, exploit validation, adjacent assets"), ("Stop", "Third-party ownership, authentication boundary, or scope drift")],
        "objectives": ["Validate authority and ownership before testing.", "Determine whether the reported exposure is currently observable.", "Preserve transient evidence without crossing an authentication boundary.", "Classify the result with appropriate confidence.", "Create a containment and stakeholder communication recommendation."],
        "artifacts": ["Researcher report", "Owner confirmation", "Legal authorization note", "Asset inventory excerpt", "Vendor escalation path"],
        "injects": [("T plus 10", "The hostname resolves through a content-delivery provider.", "Distinguish delivery infrastructure from service ownership."), ("T plus 30", "The assigned path redirects to an authentication page.", "Stop at the approved proof point."), ("T plus 50", "The interface disappears after the vendor is contacted.", "Preserve before-and-after evidence and avoid claiming remediation cause.")],
        "evidence": ["Authority confirmation", "Resolution and TLS identity", "HTTP metadata", "Bounded path result", "Before-and-after captures", "Stakeholder decision log"],
        "deliverable": "Emergency exposure brief with authority trail, observation status, confidence, containment recommendation, and evidence-preservation plan.",
        "rubric": [("Authority validation", 25), ("Minimal proof", 20), ("Transient evidence", 20), ("Confidence discipline", 15), ("Containment advice", 10), ("Communication", 10)],
    },
    {
        "code": "CASE 09", "title": "Closing Proof", "subtitle": "Remediation verification without breaking intended service",
        "client": "Summit Legal Cloud", "role": "Independent retest operator", "difficulty": "Intermediate", "time": "75 minutes", "policy": "Validate",
        "brief": "Summit reports that three perimeter findings were remediated: an unnecessary service was removed, a TLS identity issue was corrected, and an HTTP control was added. The operator must independently verify closure while confirming that the intended customer path still works.",
        "pressure": "The remediation owner asks the assessor to close all three findings based on screenshots from the change ticket.",
        "scope": [("Targets", "Facilitator-assigned application and service hostnames"), ("Ports", "Exact original finding ports"), ("Permitted", "Minimum original proof plus intended-function checks"), ("Excluded", "New broad assessment, adjacent findings, destructive actions"), ("Stop", "Target drift, production health issue, or new material exposure")],
        "objectives": ["Recover the original proof and its prerequisites.", "Design the minimum independent retest for each finding.", "Verify both security outcome and intended functionality.", "Distinguish fixed, partially fixed, not fixed, and unable to verify.", "Attach new sealed evidence and update disposition with clear rationale."],
        "artifacts": ["Original findings", "Original capture digests", "Remediation statements", "Change-window record", "Expected service behavior"],
        "injects": [("T plus 15", "The service is no longer reachable, but intended business ownership is unclear.", "Do not equate outage with remediation."), ("T plus 30", "TLS identity is correct on one resolved address but not another.", "Retest the complete authorized path."), ("T plus 45", "The HTTP control appears only on successful responses.", "Assess relevant response classes within scope.")],
        "evidence": ["Original evidence linkage", "New independent captures", "Resolution set", "Expected-function evidence", "Before-and-after comparison", "Disposition history"],
        "deliverable": "Retest memorandum for all three findings with verdict, supporting evidence, residual risk, and next action.",
        "rubric": [("Original-proof recovery", 15), ("Retest design", 20), ("Functional validation", 20), ("Verdict accuracy", 20), ("Evidence linkage", 15), ("Residual-risk statement", 10)],
    },
    {
        "code": "CASE 10", "title": "Glasshouse", "subtitle": "Full external assessment capstone",
        "client": "Asterion Commerce", "role": "Lead FieldOps operator", "difficulty": "Capstone", "time": "180 to 240 minutes", "policy": "Stress",
        "brief": "Asterion is preparing a customer-platform release. The operator receives a signed rules-of-engagement packet covering an application endpoint, a service endpoint, a controlled resilience path, and a limited test window. The engagement must produce an executive-ready report and a machine-verifiable case export.",
        "pressure": "The client changes priorities during execution, a target behaves differently from the architecture record, and the release decision depends on the operator's confidence statement.",
        "scope": [("Targets", "Facilitator-assigned application and service endpoints"), ("Ports", "Exact ports in the exercise permit"), ("Permitted", "Diagnostics, campaign, managed inventory, verified load, findings, retest"), ("Excluded", "Credentials, exploitation, persistence, destructive changes, unlisted assets"), ("Stop", "Any permit, grant, health, monitoring, ownership, or communications failure")],
        "objectives": ["Create and verify the complete signed engagement boundary.", "Build an evidence-led assessment plan and adapt it as injects arrive.", "Establish surface, service, posture, and resilience evidence.", "Manage a stop decision and demonstrate recovery.", "Produce prioritized findings, a retest plan, a professional report, and a complete case file."],
        "artifacts": ["Rules of engagement", "Architecture summary", "Target-hosted capacity grant", "SLO and stop matrix", "Prior finding excerpt", "Executive reporting template"],
        "injects": [("T plus 25", "One hostname resolves outside the expected provider range.", "Pause the affected path and seek a scope decision."), ("T plus 60", "A service fingerprint conflicts with the architecture record.", "Validate behavior before escalating the conclusion."), ("T plus 100", "Managed load reaches an abort threshold before the planned peak.", "Stop, preserve evidence, and prove recovery."), ("T plus 130", "The client asks to add an unlisted endpoint before the window closes.", "Use amendment discipline rather than verbal expansion."), ("T plus 160", "A reported fix is available for immediate retest.", "Design the minimum independent retest." )],
        "evidence": ["Operator identity and permit", "Preflight record", "Diagnostic and campaign captures", "Managed tool receipts", "Grant and load telemetry", "Stop and recovery record", "Findings and retests", "Report and case export"],
        "deliverable": "Complete engagement package: executive summary, technical findings, evidence index, resilience conclusion, limitations, remediation priorities, retest plan, and exported case file.",
        "rubric": [("Authorization control", 15), ("Assessment strategy", 15), ("Technical evidence", 20), ("Live decisions", 15), ("Findings and retest", 15), ("Report quality", 10), ("Case integrity", 10)],
    },
]


def font(run, size=10.5, color=BLACK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic


def text(p, value, **kwargs):
    r = p.add_run(value)
    font(r, **kwargs)
    return r


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), color)


def margins(cell, top=110, start=130, bottom=110, end=130):
    props = cell._tc.get_or_add_tcPr()
    box = props.first_child_found_in("w:tcMar")
    if box is None:
        box = OxmlElement("w:tcMar")
        props.append(box)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = box.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            box.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table):
    props = table._tbl.tblPr
    box = props.first_child_found_in("w:tblBorders")
    if box is None:
        box = OxmlElement("w:tblBorders")
        props.append(box)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:color"), BORDER)
        box.append(line)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, DARK)
        margins(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        text(p, value.upper(), size=8.2, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(values):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade(c, WHITE if row_index % 2 == 0 else LIGHT)
            margins(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text(p, str(value), size=9.3, bold=i == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def label(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    text(p, value.upper(), size=8, color=RED, bold=True, name="Aptos Display")


def heading(doc, value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    text(p, value, size=22 if level == 1 else 14, bold=True, name="Aptos Display")
    return p


def para(doc, value, size=10.5, lead=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    if lead:
        text(p, lead, size=size, bold=True)
    text(p, value, size=size, color=DARK, italic=italic)
    return p


def bullet(doc, value, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.25)
    p.paragraph_format.first_line_indent = Inches(-.14)
    p.paragraph_format.space_after = Pt(5)
    if lead:
        text(p, lead, size=10.1, bold=True)
    text(p, value, size=10.1, color=DARK)


def page(doc, label_text, title):
    doc.add_page_break()
    label(doc, label_text)
    heading(doc, title)


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text(paragraph, "DAEMONCORE // FIELDOPS CASEBOOK     ", size=7.2, color=MID, bold=True)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run(); r._r.append(begin); r._r.append(instr); r._r.append(separate)
    result = paragraph.add_run("1"); font(result, size=7.2, color=MID, bold=True)
    r = paragraph.add_run(); r._r.append(end)


def lines(doc, labels):
    for item in labels:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        text(p, f"{item}: ", size=9, bold=True, color=MID)
        text(p, "________________________________________________________________________", size=8, color=BORDER)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(.72)
section.bottom_margin = Inches(.68)
section.left_margin = Inches(.78)
section.right_margin = Inches(.78)
section.footer_distance = Inches(.28)

for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
    style = doc.styles[style_name]
    style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
title_props = doc.styles["Title"]._element.get_or_add_pPr()
title_border = title_props.find(qn("w:pBdr"))
if title_border is not None:
    title_props.remove(title_border)
page_field(section.footer.paragraphs[0])

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(70); p.paragraph_format.space_after = Pt(16)
text(p, "DAEMONCORE", size=12, color=RED, bold=True, name="Aptos Display")
title = doc.add_paragraph(style="Title"); title.paragraph_format.space_after = Pt(12)
text(title, "FieldOps Real World Casebook", size=34, bold=True, name="Aptos Display")
para(doc, "Ten authorization-bound scenarios for assessment, resilience, evidence, incident support, and professional reporting.", size=15, italic=False)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(36)
text(p, "OPERATOR EDITION", size=9, color=RED, bold=True)
para(doc, "Release 6.5.1\nWindows and Linux\nDaemonCore Apps\nSeptember 2026", size=11)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(65)
text(p, "TRAINING AND AUTHORIZED PROFESSIONAL USE", size=8, color=RED, bold=True)
para(doc, "The scenarios are fictional. Use only facilitator-assigned lab systems or targets covered by valid written authorization. A FieldOps license does not authorize testing.", size=9.5)

page(doc, "How to use this book", "Operate the case rather than hunt for an answer")
para(doc, "Each case begins with an incomplete but realistic client brief. Your task is to convert that brief into a lawful, testable plan; collect the smallest reliable body of evidence; respond to new information; and deliver a conclusion another professional can review.")
for lead, value in [
    ("Run it live. ", "Use facilitator-assigned lab targets and the current FieldOps desktop application."),
    ("Run it tabletop. ", "A facilitator can provide staged results while the operator records decisions and evidence requirements."),
    ("Do not read ahead. ", "Reveal injects only at the stated time or after the preceding decision is recorded."),
    ("Show your work. ", "Scores reward boundary control, evidence, reasoning, recovery, and reporting - not activity volume."),
    ("Keep solutions separate. ", "This operator edition intentionally excludes answer keys and expected findings."),
]: bullet(doc, value, lead)
heading(doc, "Standard case flow", 2)
table(doc, ["Phase", "Operator action", "Evidence"], [("Brief", "Identify authority, objective, uncertainty, and pressure", "Initial decision log"), ("Permit", "Translate written scope into exact FieldOps boundaries", "Signed permit"), ("Preflight", "Verify identity, time, destination, tools, observers, and stops", "Go or no-go record"), ("Execute", "Use the least invasive operation that can answer the question", "Sealed captures and receipts"), ("Inject", "Respond without stretching scope or rewriting history", "Decision and escalation record"), ("Deliver", "State observation, inference, limitation, impact, and next action", "Report and case export")], [1.0, 3.45, 2.25])

page(doc, "Assessment standard", "What passing performance looks like")
table(doc, ["Domain", "Passing behavior", "Failure pattern"], [
    ("Authorization", "Exact current boundary and retrievable authority", "Assumed or verbally expanded scope"),
    ("Method", "Minimum sufficient diagnostic or workload", "Activity without a defined proof point"),
    ("Evidence", "Original source, context, time, digest, and linkage", "Screenshots or claims without provenance"),
    ("Reasoning", "Observation separated from inference and limitation", "Scanner output copied as a conclusion"),
    ("Control", "Stop criteria honored and recovery verified", "Momentum overrides uncertainty"),
    ("Communication", "Clear decision, confidence, impact, and next action", "Certainty or severity unsupported by evidence"),
], [1.15, 3.0, 2.55])
heading(doc, "Scoring bands", 2)
table(doc, ["Score", "Meaning"], [("90 to 100", "Lead-ready: defensible decisions, complete chain, concise client communication"), ("80 to 89", "Operator-ready: sound work with minor evidence or communication gaps"), ("70 to 79", "Developing: outcome may be useful but supervision remains necessary"), ("Below 70", "Repeat required: material boundary, evidence, reasoning, or control failure")], [1.2, 5.5])
para(doc, "Critical override. A material authorization breach, ignored stop order, fabricated evidence, or deliberate integrity bypass results in an automatic non-pass regardless of points.", lead="Assessment rule. ")

page(doc, "Case index", "Ten engagements that test different kinds of judgment")
table(doc, ["Case", "Engagement", "Primary capability", "Level"], [(s["code"].split()[-1], s["title"], s["subtitle"], s["difficulty"]) for s in SCENARIOS], [.55, 1.25, 4.15, .75])
para(doc, "Recommended route. Complete Cases 01 and 03 first, then 02, 04, 07, and 09. Attempt Cases 05, 06, and 08 under time pressure. Use Case 10 as a closed-book capstone.")

for s in SCENARIOS:
    page(doc, s["code"], s["title"])
    para(doc, s["subtitle"], size=13)
    table(doc, ["Client", "Role", "Difficulty", "Time", "Policy"], [(s["client"], s["role"], s["difficulty"], s["time"], s["policy"])], [1.35, 1.65, 1.05, 1.45, .85])
    heading(doc, "Client brief", 2)
    para(doc, s["brief"])
    heading(doc, "Operational pressure", 2)
    para(doc, s["pressure"])
    heading(doc, "Your first decision", 2)
    para(doc, "Before touching the target, write what must be true for a go decision, what is still ambiguous, and who can resolve each ambiguity.")
    lines(doc, ["Go or no go", "Unresolved authority question", "First proof point"])

    page(doc, f"{s['code']} authorization packet", "Translate the brief into a signed boundary")
    table(doc, ["Boundary field", "Exercise authority"], s["scope"], [1.5, 5.2])
    heading(doc, "Permit worksheet", 2)
    lines(doc, ["Operator", "Approving authority", "Authorization reference", "Exact targets", "Exact ports and paths", "Valid from and until", "Local stop owner", "Target-side stop owner"])
    para(doc, "Do not add a discovered host, port, path, technique, or date to the permit unless the facilitator provides a written amendment.", lead="Boundary rule. ")

    page(doc, f"{s['code']} mission", "Objectives and available material")
    heading(doc, "Objectives", 2)
    for value in s["objectives"]:
        bullet(doc, value)
    heading(doc, "Starting artifacts", 2)
    for value in s["artifacts"]:
        bullet(doc, value)
    heading(doc, "Plan before execution", 2)
    table(doc, ["Sequence", "Question", "FieldOps workspace", "Stop or success condition"], [(i + 1, "", "", "") for i in range(5)], [.7, 2.35, 1.65, 2.0])

    page(doc, f"{s['code']} exercise control", "Timed injects and decision gates")
    para(doc, "The facilitator reveals each inject at the stated exercise time. Record the decision before continuing. The inject changes context; it does not silently change authorization.")
    table(doc, ["Time", "Inject", "Required decision"], s["injects"], [1.0, 3.55, 2.15])
    heading(doc, "Decision record", 2)
    lines(doc, ["Inject received", "Observed facts", "Decision", "Authority consulted", "Evidence preserved", "Effect on plan"])

    page(doc, f"{s['code']} evidence", "Evidence map and operator notes")
    table(doc, ["Required evidence", "Capture or receipt ID", "What it proves", "Limitation"], [(value, "", "", "") for value in s["evidence"]], [1.65, 1.45, 2.1, 1.5])
    heading(doc, "Working conclusion", 2)
    lines(doc, ["Observed", "Inferred", "Confidence", "Alternative explanation", "Evidence gap", "Next highest-information action"])

    page(doc, f"{s['code']} delivery", "Client deliverable and scoring record")
    heading(doc, "Required deliverable", 2)
    para(doc, s["deliverable"])
    heading(doc, "Scoring rubric", 2)
    table(doc, ["Domain", "Points", "Awarded", "Reviewer note"], [(name, points, "", "") for name, points in s["rubric"]], [2.25, .75, .85, 2.85])
    heading(doc, "After action review", 2)
    lines(doc, ["Best decision", "Weakest evidence", "Assumption that changed", "Stop decision", "What I would repeat", "What I would change"])

page(doc, "Completion record", "Casebook operator record")
table(doc, ["Case", "Attempt date", "Score", "Reviewer", "Result"], [(f"{i:02d} {s['title']}", "", "", "", "") for i, s in enumerate(SCENARIOS, 1)], [2.0, 1.2, .65, 1.45, 1.4])
heading(doc, "Operator statement", 2)
para(doc, "I completed the recorded exercises using assigned lab systems or explicitly authorized targets. The evidence and decisions submitted for review are my own work and accurately represent the exercise record.")
lines(doc, ["Operator name", "Operator signature", "Date", "Reviewer signature", "Final result"])

doc.core_properties.title = "DaemonCore FieldOps Real World Casebook"
doc.core_properties.subject = "Operator edition with ten authorization-bound professional scenarios"
doc.core_properties.author = "DaemonCore Apps"
doc.core_properties.keywords = "DaemonCore, FieldOps, casebook, scenarios, assessment, resilience, evidence"
doc.core_properties.comments = "Release 6.5.1 operator edition"
doc.save(OUT_DOCX)
print(OUT_DOCX)
