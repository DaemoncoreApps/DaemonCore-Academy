from __future__ import annotations

import os
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GUIDE = "academy" if "--academy" in sys.argv else os.environ.get("DAEMONCORE_GUIDE", "fieldops").lower()
if GUIDE == "academy":
    SOURCE = ROOT / "docs" / "ACADEMY-MISSION-OS-GUIDE.md"
    OUTPUT = ROOT / "docs" / "manuals" / "DaemonCore-Academy-Mission-OS-Guide-v6.0.0-beta.1.docx"
    GUIDE_NAME = "ACADEMY MISSION OS GUIDE"
    COVER_TITLE = "ACADEMY"
    COVER_SUBTITLE = "MISSION OS OPERATOR GUIDE"
    COVER_KICKER = "DAEMONCORE  //  EVIDENCE-LED OPERATOR DEVELOPMENT"
    COVER_LEAD = "Diagnose the signal. Build the route. Prove the work.\nA local-first operating guide for practical cyber training."
    COVER_NOTE = "PROGRESS IS RECORDED. CAPABILITY IS PROVEN BY THE WORK."
    DOCUMENT_TITLE = "DaemonCore Academy Mission OS Operator Guide"
    DOCUMENT_SUBJECT = "Operator guidance for DaemonCore Academy Mission OS 6.0.0-beta.1"
    DOCUMENT_KEYWORDS = "DaemonCore, Academy, Mission OS, operator guide, cyber range"
    SCREENSHOT = ROOT / "docs" / "screenshots" / "command-center.png"
    SCREENSHOT_ANCHOR = "__no_academy_screenshot__"
    SCREENSHOT_ALT = "DaemonCore Academy command center and operator progress overview"
    SCREENSHOT_TITLE = "DaemonCore Academy command center"
    SCREENSHOT_CAPTION = "DaemonCore Academy command center and recorded operator progress"
else:
    SOURCE = ROOT / "docs" / "FIELDOPS-OPERATOR-MANUAL.md"
    OUTPUT = ROOT / "docs" / "manuals" / "DaemonCore-FieldOps-Operator-Manual-v6.0.0-beta.1.docx"
    GUIDE_NAME = "FIELDOPS OPERATOR MANUAL"
    COVER_TITLE = "FIELDOPS"
    COVER_SUBTITLE = "OPERATOR MANUAL"
    COVER_KICKER = "DAEMONCORE  //  AUTHORIZED ASSESSMENT CONTROL PLANE"
    COVER_LEAD = "Signed scope. Pinned targets. Sealed evidence.\nProfessional assessment operations from one local-first desktop workspace."
    COVER_NOTE = "A LICENSE UNLOCKS THE TOOL. THE ENGAGEMENT AUTHORIZES THE TARGET."
    DOCUMENT_TITLE = "DaemonCore FieldOps Operator Manual"
    DOCUMENT_SUBJECT = "Operator guidance for DaemonCore FieldOps 6.0.0-beta.1"
    DOCUMENT_KEYWORDS = "DaemonCore, FieldOps, operator manual, authorized assessment"
    SCREENSHOT = ROOT / "docs" / "screenshots" / "fieldops-pro-gate.png"
    SCREENSHOT_ANCHOR = "Move a license to another device"
    SCREENSHOT_ALT = "FieldOps Pro license and authorization gate before activation"
    SCREENSHOT_TITLE = "FieldOps Pro gate"
    SCREENSHOT_CAPTION = "FieldOps Pro gate before commercial activation"
ICON = ROOT / "build" / "icon-sizes" / "256.png"

RED = "E33E48"
DARK = "111317"
INK = "20242A"
MID = "5E6570"
LIGHT = "F1F3F5"
PALE_RED = "FCEDEF"
WHITE = "FFFFFF"
BORDER = "D8DCE2"
CONTENT_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Aptos", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    set_run_font(run, name="Aptos", size=8, color=MID)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after in (
        ("Heading 1", 18, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 11.5, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(RED if name != "Heading 3" else DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_section(section, first=False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8 if first else 0.86)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = first


def add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"DAEMONCORE  //  {GUIDE_NAME}  //  6.0.0-BETA.1")
    set_run_font(run, name="Aptos", size=7.5, color=MID, bold=True)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "9")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), RED)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7100, 2260])
    set_repeat_table_header(table.rows[0])
    table._tbl.tblPr.remove(table._tbl.tblPr.find(qn("w:tblBorders"))) if table._tbl.tblPr.find(qn("w:tblBorders")) is not None else None
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("PUBLIC BETA  |  Copyright 2026 DaemonCore Apps")
    set_run_font(run, size=7.5, color=MID)
    right = table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    add_page_number(right)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, first=True)

    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top.paragraph_format.space_before = Pt(24)
    top.paragraph_format.space_after = Pt(26)
    icon = top.add_run().add_picture(str(ICON), width=Inches(1.35))
    icon._inline.docPr.set("descr", "DaemonCore application mark")
    icon._inline.docPr.set("title", "DaemonCore")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run(COVER_KICKER)
    set_run_font(run, name="Aptos", size=9, color=RED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(COVER_TITLE)
    set_run_font(run, name="Aptos Display", size=34, color=DARK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(COVER_SUBTITLE)
    set_run_font(run, name="Aptos", size=15, color=RED, bold=True)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.6)
    lead.paragraph_format.right_indent = Inches(0.6)
    lead.paragraph_format.space_after = Pt(34)
    run = lead.add_run(COVER_LEAD)
    set_run_font(run, size=12, color=MID)

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    set_repeat_table_header(table.rows[0])
    labels = (("PRODUCT", "DaemonCore Academy"), ("RELEASE", "6.0 Beta 1"), ("EDITION", "Public beta"))
    for index, (label, value) in enumerate(labels):
        set_cell_shading(table.cell(0, index), DARK)
        set_cell_shading(table.cell(1, index), LIGHT)
        p = table.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=7.5, color=WHITE, bold=True)
        p = table.cell(1, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9, color=INK, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(34)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(COVER_NOTE)
    set_run_font(run, size=8.5, color=RED, bold=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(8)
    run = date.add_run("30 AUGUST 2026  //  DAEMONCORE APPS")
    set_run_font(run, size=8, color=MID)

    doc.add_page_break()


def add_callout(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("OPERATOR NOTE  //  ")
    set_run_font(run, size=8.5, color=RED, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(doc, rows: list[list[str]]) -> None:
    columns = len(rows[0])
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [2100, 4100, 3160]
    else:
        widths = [CONTENT_DXA // columns] * columns
        widths[-1] += CONTENT_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, DARK)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.clear()
            run = p.add_run(value)
            set_run_font(run, size=8.4 if columns >= 3 else 9, color=WHITE if row_index == 0 else INK, bold=row_index == 0)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Cascadia Mono", size=9, color=RED)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def parse_table(lines: list[str]) -> list[list[str]]:
    parsed = []
    for line in lines:
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        parsed.append(values)
    return parsed


def render_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    seen_first_heading = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("# "):
            if not seen_first_heading:
                seen_first_heading = True
                index += 1
                continue
            doc.add_page_break()
            p = doc.add_paragraph(stripped[2:], style="Heading 1")
            p.paragraph_format.space_before = Pt(4)
            index += 1
            continue
        if stripped.startswith("## "):
            # Skip the cover subtitle already rendered on page one.
            if stripped in {"## Operator Manual", "## Mission OS Operator Guide"} and index < 4:
                index += 1
                continue
            doc.add_paragraph(stripped[3:], style="Heading 2")
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
            index += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue
        if re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, numbered.group(2))
            index += 1
            continue
        # Cover metadata is already represented in the custom cover.
        if index < 12 and (stripped.startswith("Version 6.0.0-beta.1") or stripped in {"Public beta edition", "30 August 2026"}):
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("#", ">", "|", "- ")) or re.match(r"^\d+\.\s+", next_line) or next_line == "---":
                break
            paragraph_lines.append(next_line)
            index += 1
        p = doc.add_paragraph()
        add_inline_markdown(p, " ".join(paragraph_lines))


def insert_guide_screenshot(doc: Document) -> None:
    # Keep the explanatory image adjacent to the workflow it supports.
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == SCREENSHOT_ANCHOR:
            target = paragraph
            break
    if target is None or not SCREENSHOT.exists():
        return
    image_p = OxmlElement("w:p")
    target._p.addprevious(image_p)
    image_paragraph = next(p for p in doc.paragraphs if p._p is image_p)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(8)
    image_paragraph.paragraph_format.space_after = Pt(3)
    screenshot = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.2))
    screenshot._inline.docPr.set("descr", SCREENSHOT_ALT)
    screenshot._inline.docPr.set("title", SCREENSHOT_TITLE)
    caption_p = OxmlElement("w:p")
    target._p.addprevious(caption_p)
    caption = next(p for p in doc.paragraphs if p._p is caption_p)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(SCREENSHOT_CAPTION)
    set_run_font(run, size=8, color=MID, italic=True)


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = DOCUMENT_TITLE
    props.subject = DOCUMENT_SUBJECT
    props.author = "DaemonCore Apps"
    props.keywords = DOCUMENT_KEYWORDS
    props.comments = "Public beta edition"


def main() -> int:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    qa_dir = ROOT / "tmp" / f"{GUIDE}-manual"
    if qa_dir.exists():
        shutil.rmtree(qa_dir)
    qa_dir.mkdir(parents=True)

    doc = Document()
    configure_styles(doc)
    set_document_properties(doc)
    add_cover(doc)
    for section_index, section in enumerate(doc.sections):
        configure_section(section, first=section_index == 0)
        add_running_furniture(section)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    insert_guide_screenshot(doc)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
