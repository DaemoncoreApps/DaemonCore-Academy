from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FIELDOPS-CAPABILITY-GUIDE.md"
OUT = ROOT / "output" / "guides"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "DaemonCore-FieldOps-Capability-and-Technology-Guide-v6.5.1.docx"

RED = "D71920"
BLACK = "101010"
DARK = "252525"
MID = "646464"
LIGHT = "F3F3F3"
BORDER = "D9D9D9"
WHITE = "FFFFFF"


def set_font(run, name="Aptos", size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_runs(paragraph, text, size=10.5, color=DARK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=BLACK)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Cascadia Mono", size=size - .8, color=BLACK)
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("](")]
            run = paragraph.add_run(label)
            set_font(run, size=size, bold=True, color=RED)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), color)


def margins(cell, value=110):
    props = cell._tc.get_or_add_tcPr()
    tc_mar = props.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        props.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table):
    props = table._tbl.tblPr
    group = props.first_child_found_in("w:tblBorders")
    if group is None:
        group = OxmlElement("w:tblBorders")
        props.append(group)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:color"), BORDER)
        group.append(line)


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("FIELDOPS 6.5.1     ")
    set_font(prefix, size=7.5, bold=True, color=MID)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.extend([begin, instr, separate])
    result = paragraph.add_run("1")
    set_font(result, size=7.5, bold=True, color=MID)
    paragraph.add_run()._r.append(end)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)
    repeat_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, DARK)
        margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value.upper())
        set_font(run, size=8.1, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_runs(paragraph, value, size=8.6, color=DARK)
            if index == 0:
                for item in paragraph.runs:
                    item.font.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(.7)
    section.bottom_margin = Inches(.65)
    section.left_margin = Inches(.72)
    section.right_margin = Inches(.72)
    section.header_distance = Inches(.3)
    section.footer_distance = Inches(.28)
    page_field(section.footer.paragraphs[0])
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title_props = styles["Title"]._element.get_or_add_pPr()
    border = title_props.find(qn("w:pBdr"))
    if border is not None:
        title_props.remove(border)


def cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(34)
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run("DAEMONCORE")
    set_font(run, name="Aptos Display", size=12, bold=True, color=RED)
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("FieldOps Capability and Technology Guide")
    set_font(run, name="Aptos Display", size=34, bold=True, color=BLACK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(26)
    add_runs(subtitle, "Professional assessment operations for authorized targets", size=15, color=MID)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(38)
    add_runs(meta, "RELEASE 6.5.1   WINDOWS AND LINUX   SEPTEMBER 2026", size=9, color=RED)
    intro = doc.add_paragraph()
    intro.paragraph_format.line_spacing = 1.18
    intro.paragraph_format.space_after = Pt(14)
    add_runs(intro, "A complete buyer and operator reference to FieldOps diagnostics, managed tools, resilience testing, evidence controls, findings, reporting, supported technologies, deployment requirements and current boundaries.", size=12, color=DARK)
    notice = doc.add_paragraph()
    notice.paragraph_format.space_before = Pt(54)
    notice.paragraph_format.line_spacing = 1.12
    add_runs(notice, "CONTROLLED PROFESSIONAL USE. A software license does not authorize a target. Operate only under valid written authorization, exact scope, an active testing window and an agreed stop procedure.", size=9.5, color=BLACK)
    doc.add_page_break()


def render_markdown(doc, lines):
    index = 0
    first_h1 = True
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            first_h1 = False
            index += 1
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line[3:])
            set_font(run, name="Aptos Display", size=20, bold=True, color=BLACK)
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(line[4:])
            set_font(run, name="Aptos Display", size=13, bold=True, color=BLACK)
            index += 1
            continue
        if line.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [[cell.strip() for cell in item.strip("|").split("|")] for item in table_lines]
            headers = parsed[0]
            rows = [row for row in parsed[2:] if len(row) == len(headers)]
            add_table(doc, headers, rows)
            continue
        if re.match(r"^- ", line):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(.25)
            paragraph.paragraph_format.first_line_indent = Inches(-.15)
            paragraph.paragraph_format.space_after = Pt(3)
            add_runs(paragraph, line[2:], size=10, color=DARK)
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(.25)
            paragraph.paragraph_format.first_line_indent = Inches(-.15)
            paragraph.paragraph_format.space_after = Pt(3)
            add_runs(paragraph, re.sub(r"^\d+\. ", "", line), size=10, color=DARK)
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.12
        add_runs(paragraph, line, size=10.5, color=DARK)
        index += 1


doc = Document()
configure(doc)
cover(doc)
source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
render_markdown(doc, source_lines)

doc.core_properties.title = "DaemonCore FieldOps Capability and Technology Guide"
doc.core_properties.subject = "FieldOps release 6.5.1 capabilities, technologies and operating requirements"
doc.core_properties.author = "DaemonCore Apps"
doc.core_properties.keywords = "DaemonCore, FieldOps, capability, technology, pentesting, Nmap, k6, evidence"
doc.core_properties.comments = "Release 6.5.1"
doc.save(DOCX)
print(DOCX)
